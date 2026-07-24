from __future__ import annotations

import re
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from threading import local

import requests
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


DOCS_DIR = Path(__file__).resolve().parent
SOURCE_PATH = DOCS_DIR / "UEditor_Huong_dan_su_dung_chi_tiet_VI.docx"
OUTPUT_PATH = DOCS_DIR / "UEditor_Detailed_User_Guide_EN.docx"
TRANSLATE_URL = "https://translate.googleapis.com/translate_a/single"
TYPEFACE = "Times New Roman"
THREAD_STATE = local()


MANUAL_TRANSLATIONS = {
    "UNDERVERSE  •  TÀI LIỆU HƯỚNG DẪN": "UNDERVERSE • USER DOCUMENTATION",
    "HƯỚNG DẪN SỬ DỤNG": "USER GUIDE",
    "HƯỚNG DẪN SỬ DỤNG\nUEDITOR": "USER GUIDE\nUEDITOR",
    "Tài liệu hướng dẫn chi tiết dành cho người dùng cuối": "Detailed Guide for End Users",
    "Soạn thảo văn bản • Định dạng • Hình ảnh • Liên kết • Khối nội dung • Bảng nâng cao • Công thức • Phím tắt • Xử lý lỗi": (
        "Text Editing • Formatting • Images • Links • Content Blocks • Advanced Tables • "
        "Formulas • Shortcuts • Troubleshooting"
    ),
    "HƯỚNG DẪN SỬ DỤNG  •  PHIÊN BẢN 2.0": "USER GUIDE • VERSION 2.0",
    "THÔNG TIN — Ngôn ngữ của tài liệu": "INFORMATION — Document Language",
    "Phần giải thích được viết bằng tiếng Việt. Theo yêu cầu, toàn bộ ảnh chụp minh họa dùng giao diện tiếng Anh. Tên tiếng Anh quan trọng được đặt trong ngoặc để người đọc dễ đối chiếu.": (
        "This edition is written in English, and all screenshots use the English interface. "
        "Section 19 provides an English–Vietnamese terminology reference."
    ),
    "Phiên bản tài liệu": "Document Version",
    "Ngày cập nhật": "Last Updated",
    "Đối tượng": "Audience",
    "Khách hàng / người dùng cuối": "Customers / End Users",
    "Phạm vi": "Scope",
    "Thông tin và phạm vi tài liệu": "Document Information and Scope",
    "Cách đọc tên chức năng": "How Feature Names Are Presented",
    "Tên tiếng Việt mô tả ý nghĩa sử dụng; tên tiếng Anh trong ngoặc là nhãn xuất hiện trên ảnh minh họa.": (
        "Feature names are shown in English to match the screenshots. Vietnamese equivalents are listed in Section 19."
    ),
    "LƯU Ý — Giao diện có thể khác đôi chút": "NOTE — The Interface May Vary",
    "UEditor cho phép hệ thống bật/tắt thanh công cụ, menu cổ điển, menu nổi, bộ đếm và một số chức năng tích hợp. Vì vậy, một màn hình nghiệp vụ có thể không hiển thị toàn bộ nút được mô tả trong tài liệu.": (
        "UEditor allows each application screen to enable or disable toolbars, classic menus, "
        "bubble menus, counters, and selected integrated features. Therefore, a screen may not "
        "display every button described in this guide."
    ),
    "Mục lục": "Table of Contents",
    "Nhu cầu": "Task",
    "Quy ước mức độ đảm bảo của phím tắt": "Shortcut Reliability Levels",
    "1. Bắt đầu nhanh trong 3 phút": "1. Quick Start: 3 Minutes",
    "1.1 Quy trình soạn và lưu nội dung": "1.1 Creating and Saving Content",
    "1.2 Bảng tra cứu nhanh": "1.2 Quick Reference",
    "2. Làm quen với giao diện": "2. Getting Familiar with the Interface",
    "2.1 Các vùng chính": "2.1 Main Interface Areas",
    "2.3 Chế độ chỉ đọc": "2.3 Read-Only Mode",
    "3. Soạn thảo cơ bản": "3. Basic Editing",
    "3.1 Nhập, chọn và xóa nội dung": "3.1 Creating, Selecting, and Deleting Content",
    "3.2 Cắt, sao chép và dán": "3.2 Cut, Copy, and Paste",
    "3.3 Thụt lề bằng Tab": "3.3 Indenting with Tab",
    "4. Định dạng văn bản và đoạn": "4. Text and Paragraph Formatting",
    "4.1 Chọn kiểu đoạn": "4.1 Paragraph Styles",
    "4.2 Font, cỡ chữ, giãn dòng và giãn chữ": "4.2 Font, Font Size, Line Height, and Letter Spacing",
    "4.3 Định dạng ký tự": "4.3 Character Formatting",
    "4.4 Màu chữ và màu đánh dấu": "4.4 Text and Highlight Colors",
    "4.5 Căn lề, hướng chữ và ngắt dòng": "4.5 Alignment, Text Direction, and Line Breaks",
    "4.6 Xóa định dạng": "4.6 Clear Formatting",
    "5. Danh sách, công việc và trường tương tác": "5. Lists, Tasks, and Form Fields",
    "5.2 Danh sách công việc (Task List)": "5.2 Task Lists",
    "5.3 Checkbox biểu mẫu": "5.3 Checkbox Form Fields",
    "5.4 Radio biểu mẫu": "5.4 Radio Button Form Fields",
    "6. Liên kết": "6. Links",
    "6.1 Chèn liên kết vào chữ": "6.1 Inserting Links into Text",
    "6.2 Xem, sửa và gỡ liên kết": "6.2 Opening, Editing, and Removing Links",
    "6.3 Địa chỉ hợp lệ": "6.3 Valid URLs",
    "7. Hình ảnh": "7. Images",
    "7.1 Các cách chèn ảnh": "7.1 Inserting Images",
    "7.2 Mô tả ảnh (Alt text)": "7.2 Image Descriptions (Alt Text)",
    "7.3 Chọn bố cục và kích thước": "7.3 Image Layout and Size",
    "7.4 Kéo đổi kích thước ảnh": "7.4 Resizing Images by Dragging",
    "7.5 Ảnh khi lưu": "7.5 How Images Are Handled When Saving",
    "8. Lệnh nhanh bằng dấu /": "8. Slash Commands",
    "8.1 Cách mở và điều khiển": "8.1 Opening and Navigating the Command Menu",
    "8.2 Danh sách lệnh": "8.2 Command Reference",
    "9. Emoji và các khối nội dung nâng cao": "9. Emoji and Advanced Content Blocks",
    "9.1 Chèn emoji bằng dấu :": "9.1 Inserting Emoji with :",
    "9.2 Hộp thông tin (Callout)": "9.2 Callouts",
    "9.3 Thẻ liên kết (Bookmark Card)": "9.3 Bookmark Cards",
    "9.4 Tệp đính kèm (File Attachment)": "9.4 File Attachments",
    "9.5 Khối mã (Code Block)": "9.5 Code Blocks",
    "10. Menu cổ điển": "10. Classic Menus",
    "10.6 Tools và Table": "10.6 Tools and Table",
    "THÔNG TIN — Cách hiểu con số 131": "INFORMATION — How to Interpret the 131-item Count",
    "THÔNG TIN — Phân biệt bôi đen chữ và chọn ô": "INFORMATION — Selecting Text vs. Selecting a Cell",
    "2.2 Trạng thái nút": "2.2 Button States",
    "5.1 Danh sách dấu chấm và danh sách số": "5.1 Bulleted and Numbered Lists",
    "9.6 Đường phân cách": "9.6 Horizontal Rule",
    "11. Bảng: hướng dẫn đầy đủ": "11. Tables: Complete User Guide",
    "11.1 Nhận biết các vùng điều khiển bảng": "11.1 Table Controls",
    "11.2 Chèn bảng mới": "11.2 Inserting a New Table",
    "11.3 Dán bảng từ Excel, Google Sheets hoặc trang web": (
        "11.3 Pasting a Table from Excel, Google Sheets, or a Website"
    ),
    "11.4 Chọn ô và di chuyển bằng bàn phím": "11.4 Selecting Cells and Navigating with the Keyboard",
    "11.5 Menu ba chấm: lệnh áp dụng cho bảng": "11.5 Three-Dot Menu: Whole-Table Commands",
    "11.6 Menu hàng": "11.6 Row Menu",
    "11.7 Menu cột": "11.7 Column Menu",
    "11.8 Thêm nhanh nhiều hàng hoặc cột": "11.8 Adding Multiple Rows or Columns Quickly",
    "11.9 Đổi kích thước cột, hàng và toàn bảng": "11.9 Resizing Columns, Rows, and the Entire Table",
    "11.10 Hàng/cột tiêu đề và vị trí toàn bảng": "11.10 Header Rows, Header Columns, and Table Alignment",
    "11.11 Gộp và tách ô": "11.11 Merging and Splitting Cells",
    "11.12 Mở thanh định dạng ô": "11.12 Opening the Cell Formatting Toolbar",
    "11.13 Màu nền và căn chữ theo chiều ngang": "11.13 Cell Background and Horizontal Alignment",
    "11.14 Căn dọc, hướng chữ và xuống dòng": "11.14 Vertical Alignment, Text Direction, and Wrapping",
    "11.15 Định dạng viền ô": "11.15 Cell Border Formatting",
    "11.16 Công thức và định dạng số tại ô": "11.16 Cell Formulas and Number Formats",
    "11.17 Phân biệt Clear, Delete và Convert to Value": (
        "11.17 Clear vs. Delete vs. Convert to Value"
    ),
    "11.18 Quy trình thao tác bảng an toàn": "11.18 Safe Table Editing Workflow",
    "11.19 Những chức năng chưa có trong phiên bản hiện tại": "11.19 Features Not Available in the Current Version",
    "11.20 Ma trận kiểm tra 131 thao tác và tùy chọn Table": "11.20 Reference Matrix of 131 Table Operations and Options",
    "12. Công thức trong bảng": "12. Table Formulas",
    "12.1 Khái niệm địa chỉ ô": "12.1 Understanding Cell Addresses",
    "12.2 Nhập công thức trực tiếp": "12.2 Entering a Formula Directly",
    "12.3 Các hàm hỗ trợ": "12.3 Supported Functions",
    "12.4 Toán tử và cú pháp": "12.4 Operators and Syntax",
    "12.5 Chọn vùng bằng chuột": "12.5 Selecting a Cell Range with the Mouse",
    "12.6 Thanh công thức": "12.6 Formula Bar",
    "12.7 Định dạng kết quả": "12.7 Result Formats",
    "12.8 Tính lại tự động và thủ công": "12.8 Automatic and Manual Recalculation",
    "12.9 Các lỗi công thức": "12.9 Formula Errors",
    "13. Phím tắt và thao tác bàn phím": "13. Keyboard Shortcuts",
    "13.1 Phím tắt định dạng được UEditor xử lý": "13.1 Shortcuts Handled by UEditor",
    "13.2 Lịch sử, chọn và clipboard": "13.2 History, Selection, and Clipboard",
    "13.3 Phím theo ngữ cảnh": "13.3 Context-Sensitive Keys",
    "13.4 Những nhãn phím tắt cần hiểu đúng": "13.4 How to Interpret Shortcut Labels",
    "14. Gõ nhanh theo cú pháp": "14. Syntax-Based Typing Shortcuts",
    "15. Lưu, xem trước và xuất nội dung": "15. Saving, Previewing, and Exporting Content",
    "15.1 Lưu nội dung": "15.1 Saving Content",
    "15.2 Xem trước": "15.2 Previewing Content",
    "15.3 Xuất HTML": "15.3 Exporting HTML",
    "15.4 Chỉnh HTML trực tiếp": "15.4 Editing HTML Source",
    "16. Thực hành tốt khi soạn nội dung": "16. Content Authoring Best Practices",
    "16.1 Cấu trúc và khả năng đọc": "16.1 Structure and Readability",
    "16.2 Khả năng tiếp cận": "16.2 Accessibility",
    "16.3 Hiệu năng và độ ổn định": "16.3 Performance and Stability",
    "17. Xử lý sự cố thường gặp": "17. Troubleshooting",
    "17.1 Quy trình phục hồi an toàn": "17.1 Safe Recovery Workflow",
    "18. Danh mục tính năng đầy đủ": "18. Complete Feature Inventory",
    "18.1 Tổng quan theo nhóm": "18.1 Feature Groups",
    "18.2 Ma trận 100 thao tác và tùy chọn ngoài Table": "18.2 Reference Matrix of 100 Non-Table Operations and Options",
    "19. Thuật ngữ đối chiếu Anh – Việt": "19. English–Vietnamese Terminology Reference",
    "Nhấp đúp chuột trái lên phần chữ": "Double-click the text with the left mouse button",
    "Nhấp chuột trái ba lần lên phần chữ": "Triple-click the text with the left mouse button",
    "Nhấp đúp chuột trái vào vùng trống trong ô": "Double-click a blank area in the cell with the left mouse button",
    "Nhấp đúp chuột trái lên chữ": "Double-click the text with the left mouse button",
    "Nhấp chuột trái ba lần lên chữ": "Triple-click the text with the left mouse button",
    "Nhấp đúp chuột trái vào vùng trống": "Double-click a blank area in the cell with the left mouse button",
    "Bôi đen nội dung chữ trong ô.": "Select the text within the cell.",
    "Bôi đen phần chữ trong ô để định dạng nội dung.": "Select the text within the cell for text-level formatting.",
    "Chọn nguyên ô thay vì chỉ bôi đen phần chữ.": "Select the entire cell rather than only the text.",
    "Chọn nguyên ô và mở thanh định dạng ô khi menu nổi được bật.": (
        "Select the entire cell and open the cell formatting toolbar when the bubble menu is enabled."
    ),
    "Chọn nguyên ô.": "Select the entire cell.",
    "Nhấp đúp chuột trái trực tiếp lên chữ để bôi đen phần chữ trong ô. Muốn chọn nguyên ô, nhấp chuột trái ba lần lên chữ hoặc nhấp đúp vào vùng trống trong ô.": (
        "Double-click directly on the text with the left mouse button to select text within the cell. "
        "To select the entire cell, triple-click the text or double-click a blank area inside the cell."
    ),
    "Mục lục sẽ được cập nhật khi mở tài liệu trong Microsoft Word.": (
        "The table of contents will update when the document is opened in Microsoft Word."
    ),
    "Giữ cấu trúc khi dán": "Preserve Structure When Pasting",
    "UEditor cố giữ gộp ô, độ rộng, chiều cao, màu, viền và định dạng chữ nếu nguồn cung cấp.": (
        "UEditor preserves merged cells, width, height, colors, borders, and text formatting when supplied by the source."
    ),
    "Mở rộng bảng trong lúc nhập.": "Expand the table while entering data.",
    "Cuộn ngang bảng rộng": "Scroll a Wide Table Horizontally",
    "Nội dung ngoài bảng.": "Content outside the table.",
    "Xem bộ đếm dưới editor khi màn hình bật tính năng.": "View the counter below the editor when enabled.",
    "Giới hạn ký tự": "Character Limit",
    "Theo dõi bộ đếm/giới hạn do màn hình cấu hình.": "Monitor the counter or limit configured for the editor.",
    "Nếu thanh không xuất hiện, dùng toolbar phía trên editor. Một số màn hình có thể tắt menu nổi.": (
        "If the toolbar does not appear, use the main toolbar above the editor. "
        "Some application screens may disable the bubble menu."
    ),
}


POLISH_REPLACEMENTS = (
    ("Business monitors", "Application screens"),
    ("business monitors", "application screens"),
    ("business monitor", "application screen"),
    ("built-in monitor", "integrated screen"),
    ("nodes described", "buttons described"),
    ("Board control", "Table controls"),
    ("board control", "table controls"),
    ("Boards larger than", "Tables larger than"),
    ("boards larger than", "tables larger than"),
    ("edge of the board", "edge of the table"),
    ("board horizontal scroll", "horizontal table scrolling"),
    ("Node status", "Button states"),
    ("Dot list and number list", "Bulleted and numbered lists"),
    ("Dot list", "Bulleted list"),
    ("dot list", "bulleted list"),
    ("List of numbers", "Numbered list"),
    ("Separation line", "Horizontal rule"),
    ("separation line", "horizontal rule"),
    ("Full-panel control menu", "Whole-table control menu"),
    ("full-panel control menu", "whole-table control menu"),
    ("Cell format bar", "Cell formatting toolbar"),
    ("cell format bar", "cell formatting toolbar"),
    ("cross button", "diagonal resize handle"),
    ("Commerge", "Merge"),
    ("when gluing", "when pasting"),
    ("cell aggregation", "merged cells"),
    ("Out-of-table content", "Content outside the table"),
    ("off-table", "non-table"),
    ("Photo/file", "Image/file"),
    ("photo/file", "image/file"),
    ("Photos", "Images"),
    ("photos", "images"),
    ("Photo", "Image"),
    ("photo", "image"),
    ("Monitor Uploading status", "Check the Uploading status"),
    ("monitor Uploading status", "check the Uploading status"),
    ("business screen", "application screen"),
    ("business screens", "application screens"),
    ("; You", ". You"),
    ("; The", ". The"),
    ("; Press", ". Press"),
)


def polish_translation(text: str) -> str:
    for source, target in POLISH_REPLACEMENTS:
        text = text.replace(source, target)
    return text.replace("\u200b", "").replace("\ufeff", "")


def iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested_table in cell.tables:
                yield from iter_table_paragraphs(nested_table)


def all_paragraphs(document: Document):
    yield from document.paragraphs
    for table in document.tables:
        yield from iter_table_paragraphs(table)
    for section in document.sections:
        for part in (section.header, section.footer):
            yield from part.paragraphs
            for table in part.tables:
                yield from iter_table_paragraphs(table)


def get_session() -> requests.Session:
    session = getattr(THREAD_STATE, "session", None)
    if session is None:
        session = requests.Session()
        session.headers.update({"User-Agent": "Mozilla/5.0"})
        THREAD_STATE.session = session
    return session


def translate_batch(batch: list[tuple[int, str]]) -> dict[int, str]:
    payload = "\n".join(f"[[[UESEG{index:05d}]]]{text}" for index, text in batch)
    response = get_session().post(
        TRANSLATE_URL,
        data={"client": "gtx", "sl": "vi", "tl": "en", "dt": "t", "q": payload},
        timeout=45,
    )
    response.raise_for_status()
    data = response.json()
    translated = "".join(part[0] for part in data[0] if part and part[0])
    pattern = re.compile(r"\[\[\[UESEG(\d{5})\]\]\](.*?)(?=\n?\[\[\[UESEG|\Z)", re.DOTALL)
    result = {int(index): value.strip() for index, value in pattern.findall(translated)}
    if len(result) != len(batch):
        missing = [index for index, _ in batch if index not in result]
        raise RuntimeError(f"Translation response omitted segments: {missing}")
    return result


def make_batches(entries: list[tuple[int, str]], max_chars: int = 3200) -> list[list[tuple[int, str]]]:
    batches: list[list[tuple[int, str]]] = []
    current: list[tuple[int, str]] = []
    current_length = 0
    for entry in entries:
        extra = len(entry[1]) + 24
        if current and current_length + extra > max_chars:
            batches.append(current)
            current = []
            current_length = 0
        current.append(entry)
        current_length += extra
    if current:
        batches.append(current)
    return batches


def translate_texts(texts: list[str]) -> dict[str, str]:
    unique = list(dict.fromkeys(text for text in texts if text.strip()))
    translated: dict[str, str] = {
        source: polish_translation(target)
        for source, target in MANUAL_TRANSLATIONS.items()
        if source in unique
    }
    pending = [(index, text) for index, text in enumerate(unique) if text not in translated]
    index_to_text = {index: text for index, text in pending}
    batches = make_batches(pending)

    with ThreadPoolExecutor(max_workers=6) as executor:
        futures = {executor.submit(translate_batch, batch): batch for batch in batches}
        for future in as_completed(futures):
            result = future.result()
            for index, target in result.items():
                translated[index_to_text[index]] = polish_translation(target)

    return translated


def paragraph_translation_source(paragraph) -> str | None:
    if not paragraph.text.strip():
        return None
    if paragraph._p.xpath(".//w:fldChar") or paragraph._p.xpath(".//w:instrText"):
        return None
    text_runs = [run for run in paragraph.runs if run.text]
    if (
        len(text_runs) == 2
        and re.fullmatch(r"\d+\.\s+", text_runs[0].text)
        and text_runs[1].text.strip()
    ):
        return text_runs[1].text
    return paragraph.text


def set_paragraph_translation(paragraph, translated: str) -> None:
    text_runs = [run for run in paragraph.runs if run.text]
    if (
        len(text_runs) == 2
        and re.fullmatch(r"\d+\.\s+", text_runs[0].text)
        and text_runs[1].text.strip()
    ):
        text_runs[1].text = translated
        return

    target_run = text_runs[0] if text_runs else paragraph.add_run()
    target_run.text = translated
    for run in text_runs[1:]:
        run.text = ""


def set_run_typeface_and_language(run) -> None:
    run.font.name = TYPEFACE
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{slot}"), TYPEFACE)
    language = r_pr.find(qn("w:lang"))
    if language is None:
        language = OxmlElement("w:lang")
        r_pr.append(language)
    for slot in ("val", "eastAsia", "bidi"):
        language.set(qn(f"w:{slot}"), "en-US")


def normalize_styles(document: Document) -> None:
    for style in document.styles:
        if not hasattr(style, "font"):
            continue
        style.font.name = TYPEFACE
        r_pr = style.element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.insert(0, r_fonts)
        for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
            r_fonts.set(qn(f"w:{slot}"), TYPEFACE)
        language = r_pr.find(qn("w:lang"))
        if language is None:
            language = OxmlElement("w:lang")
            r_pr.append(language)
        for slot in ("val", "eastAsia", "bidi"):
            language.set(qn(f"w:{slot}"), "en-US")


def main() -> None:
    if not SOURCE_PATH.exists():
        raise FileNotFoundError(f"Generate the Vietnamese guide first: {SOURCE_PATH}")

    document = Document(SOURCE_PATH)
    paragraphs = list(all_paragraphs(document))
    paragraph_sources = {
        id(paragraph): source
        for paragraph in paragraphs
        if (source := paragraph_translation_source(paragraph))
    }
    field_text_nodes = [
        node
        for paragraph in paragraphs
        if paragraph._p.xpath(".//w:fldChar") or paragraph._p.xpath(".//w:instrText")
        for node in paragraph._p.xpath(".//w:t")
        if node.text and node.text.strip()
    ]
    sources = list(paragraph_sources.values()) + [node.text for node in field_text_nodes]
    translations = translate_texts(sources)

    for paragraph in paragraphs:
        source = paragraph_sources.get(id(paragraph))
        if source:
            set_paragraph_translation(paragraph, translations[source])
    for node in field_text_nodes:
        node.text = translations[node.text]

    document.core_properties.title = "Detailed UEditor User Guide"
    document.core_properties.subject = "English end-user guide"
    document.core_properties.author = "Underverse"
    document.core_properties.keywords = "UEditor, user guide, editor, tables, formulas"
    document.core_properties.comments = "Detailed UEditor user guide."

    normalize_styles(document)
    for paragraph in paragraphs:
        for run in paragraph.runs:
            set_run_typeface_and_language(run)

    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
