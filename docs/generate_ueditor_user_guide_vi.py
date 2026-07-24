from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
DOCS_DIR = ROOT / "docs"
ASSET_DIR = DOCS_DIR / "assets" / "ueditor"
OUTPUT_PATH = DOCS_DIR / "UEditor_Huong_dan_su_dung_chi_tiet_VI.docx"

NAVY = "0F172A"
BLUE = "2563EB"
BLUE_DARK = "1D4ED8"
BLUE_LIGHT = "EFF6FF"
SKY_LIGHT = "F0F9FF"
GREEN = "15803D"
GREEN_LIGHT = "F0FDF4"
AMBER = "B45309"
AMBER_LIGHT = "FFFBEB"
RED = "B91C1C"
RED_LIGHT = "FEF2F2"
SLATE = "475569"
SLATE_LIGHT = "F8FAFC"
GRAY_BORDER = "CBD5E1"
WHITE = "FFFFFF"
BODY_FONT = "Times New Roman"
CODE_FONT = "Times New Roman"


def set_run_typeface(run, name: str = BODY_FONT) -> None:
    """Apply one typeface to every Word font slot, including Vietnamese text."""
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{slot}"), name)

    language = r_pr.find(qn("w:lang"))
    if language is None:
        language = OxmlElement("w:lang")
        r_pr.append(language)
    language.set(qn("w:val"), "vi-VN")
    language.set(qn("w:eastAsia"), "vi-VN")


def set_style_typeface(style, name: str = BODY_FONT) -> None:
    style.font.name = name
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{slot}"), name)

    language = r_pr.find(qn("w:lang"))
    if language is None:
        language = OxmlElement("w:lang")
        r_pr.append(language)
    language.set(qn("w:val"), "vi-VN")
    language.set(qn("w:eastAsia"), "vi-VN")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **edges) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in edges:
            continue
        edge_data = edges[edge]
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key in ("val", "sz", "space", "color"):
            if key in edge_data:
                element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_column_widths(table, widths_cm: list[float]) -> None:
    for row in table.rows:
        for index, width in enumerate(widths_cm):
            if index < len(row.cells):
                row.cells[index].width = Cm(width)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    fld_char_begin.set(qn("w:dirty"), "true")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fallback = OxmlElement("w:t")
    fallback.text = "Mục lục sẽ được cập nhật khi mở tài liệu trong Microsoft Word."
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_separate, fallback, fld_char_end])


def set_repeat_header_settings(document: Document) -> None:
    settings = document.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def configure_document() -> Document:
    document = Document()
    document.core_properties.title = "Hướng dẫn sử dụng UEditor chi tiết"
    document.core_properties.subject = "Tài liệu hướng dẫn người dùng cuối bằng tiếng Việt"
    document.core_properties.author = "Underverse"
    document.core_properties.keywords = "UEditor, hướng dẫn sử dụng, trình soạn thảo, bảng, công thức"
    document.core_properties.comments = "Tài liệu hướng dẫn sử dụng UEditor."

    section = document.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.7)
    section.header_distance = Cm(0.75)
    section.footer_distance = Cm(0.75)

    styles = document.styles
    normal = styles["Normal"]
    set_style_typeface(normal)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(5.5)
    normal.paragraph_format.line_spacing = 1.16
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in (
        ("Title", 28, NAVY, 0, 12),
        ("Subtitle", 15, SLATE, 0, 8),
        ("Heading 1", 19, NAVY, 16, 8),
        ("Heading 2", 14.5, BLUE_DARK, 12, 5),
        ("Heading 3", 11.5, NAVY, 8, 3),
    ):
        style = styles[name]
        set_style_typeface(style)
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    styles["Heading 1"].paragraph_format.page_break_before = False

    if "Guide Caption" not in styles:
        caption = styles.add_style("Guide Caption", WD_STYLE_TYPE.PARAGRAPH)
        set_style_typeface(caption)
        caption.font.size = Pt(9)
        caption.font.italic = True
        caption.font.color.rgb = RGBColor.from_string(SLATE)
        caption.paragraph_format.space_before = Pt(3)
        caption.paragraph_format.space_after = Pt(8)
        caption.paragraph_format.keep_with_next = False

    if "Guide Code" not in styles:
        code = styles.add_style("Guide Code", WD_STYLE_TYPE.PARAGRAPH)
        set_style_typeface(code, CODE_FONT)
        code.font.size = Pt(9.5)
        code.font.color.rgb = RGBColor.from_string(NAVY)
        code.paragraph_format.left_indent = Cm(0.35)
        code.paragraph_format.right_indent = Cm(0.35)
        code.paragraph_format.space_before = Pt(2)
        code.paragraph_format.space_after = Pt(2)

    for style_name in ("List Bullet", "List Bullet 2", "List Number", "List Number 2"):
        style = styles[style_name]
        set_style_typeface(style)
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(3.5)

    header = section.header
    header.is_linked_to_previous = False
    header_table = header.add_table(rows=1, cols=2, width=Cm(17.4))
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False
    set_column_widths(header_table, [9.8, 7.6])
    left = header_table.cell(0, 0)
    right = header_table.cell(0, 1)
    left.text = "UEDITOR  |  HƯỚNG DẪN NGƯỜI DÙNG"
    right.text = "Cập nhật 24/07/2026"
    for cell in (left, right):
        set_cell_margins(cell, top=25, bottom=35, start=0, end=0)
        set_cell_border(cell, bottom={"val": "single", "sz": "8", "color": BLUE})
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                set_run_typeface(run)
                run.font.size = Pt(8)
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(SLATE)
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer
    footer.is_linked_to_previous = False
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.add_run("Tài liệu dành cho khách hàng  •  Trang ")
    add_page_number(footer_paragraph)
    for run in footer_paragraph.runs:
        set_run_typeface(run)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(SLATE)

    set_repeat_header_settings(document)
    return document


def add_heading(document: Document, text: str, level: int = 1, page_break: bool = False):
    if page_break:
        document.add_page_break()
    return document.add_heading(text, level=level)


def add_paragraph(document: Document, text: str = "", bold_prefix: str | None = None):
    paragraph = document.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        first.bold = True
        paragraph.add_run(text[len(bold_prefix):])
    else:
        paragraph.add_run(text)
    return paragraph


def add_bullets(document: Document, items: list[str], level: int = 0) -> None:
    style = "List Bullet 2" if level else "List Bullet"
    for item in items:
        paragraph = document.add_paragraph(style=style)
        paragraph.add_run(item)


def add_numbered_steps(document: Document, items: list[str], level: int = 0) -> None:
    left_indent = 1.15 if level else 0.75
    for index, item in enumerate(items, start=1):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(left_indent)
        paragraph.paragraph_format.first_line_indent = Cm(-0.55)
        paragraph.paragraph_format.space_after = Pt(3.5)
        marker = paragraph.add_run(f"{index}. ")
        marker.bold = True
        marker.font.color.rgb = RGBColor.from_string(BLUE_DARK)
        paragraph.add_run(item)


def add_code(document: Document, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, SLATE_LIGHT)
    set_cell_margins(cell, top=120, start=150, bottom=120, end=150)
    set_cell_border(
        cell,
        top={"val": "single", "sz": "6", "color": GRAY_BORDER},
        left={"val": "single", "sz": "6", "color": GRAY_BORDER},
        bottom={"val": "single", "sz": "6", "color": GRAY_BORDER},
        right={"val": "single", "sz": "6", "color": GRAY_BORDER},
    )
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.style = document.styles["Guide Code"]
    paragraph.add_run(text)


def add_callout(document: Document, kind: str, title: str, text: str) -> None:
    palette = {
        "info": (BLUE, BLUE_LIGHT, "THÔNG TIN"),
        "tip": (GREEN, GREEN_LIGHT, "MẸO"),
        "warning": (AMBER, AMBER_LIGHT, "LƯU Ý"),
        "danger": (RED, RED_LIGHT, "CẢNH BÁO"),
    }
    color, fill, default_label = palette[kind]
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, start=165, bottom=120, end=165)
    set_cell_border(cell, left={"val": "single", "sz": "28", "color": color})
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    label = paragraph.add_run(f"{default_label} — {title}")
    label.bold = True
    label.font.size = Pt(9.5)
    label.font.color.rgb = RGBColor.from_string(color)
    body = cell.add_paragraph(text)
    body.paragraph_format.space_after = Pt(0)
    body.paragraph_format.line_spacing = 1.1


def add_table(
    document: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_cm: list[float] | None = None,
    compact: bool = False,
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = widths_cm is None
    table.style = "Table Grid"
    header_row = table.rows[0]
    repeat_table_header(header_row)
    for index, header in enumerate(headers):
        cell = header_row.cells[index]
        cell.text = header
        set_cell_shading(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, top=85 if compact else 110, start=90, bottom=85 if compact else 110, end=90)
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                set_run_typeface(run)
                run.font.size = Pt(8.7 if compact else 9.2)
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(WHITE)

    for row_index, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        for col_index, value in enumerate(values):
            cell = row.cells[col_index]
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell, top=70 if compact else 95, start=90, bottom=70 if compact else 95, end=90)
            if row_index % 2 == 1:
                set_cell_shading(cell, SLATE_LIGHT)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.06 if compact else 1.1
                for run in paragraph.runs:
                    set_run_typeface(run)
                    run.font.size = Pt(8.5 if compact else 9.2)
                    run.font.color.rgb = RGBColor.from_string(NAVY)
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": GRAY_BORDER},
                left={"val": "single", "sz": "4", "color": GRAY_BORDER},
                bottom={"val": "single", "sz": "4", "color": GRAY_BORDER},
                right={"val": "single", "sz": "4", "color": GRAY_BORDER},
            )

    if widths_cm:
        set_column_widths(table, widths_cm)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_image(document: Document, filename: str, caption: str, width_cm: float = 16.5) -> None:
    path = ASSET_DIR / filename
    if not path.exists():
        add_callout(document, "warning", "Thiếu ảnh minh họa", f"Không tìm thấy tệp ảnh: {path.name}")
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    caption_paragraph = document.add_paragraph(style="Guide Caption")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.add_run(caption)


def get_table_feature_rows() -> list[list[str]]:
    """User-visible Table operations/options, kept as a coverage checklist."""
    items = [
        ("Chèn và nhập", "Chèn bằng lưới trên toolbar", "Bấm Insert Table, rê chọn từ 1×1 đến 8×8 rồi bấm."),
        ("Chèn và nhập", "Chèn từ menu Insert", "Chọn Insert → Table, sau đó chọn số hàng và cột trên lưới."),
        ("Chèn và nhập", "Chèn từ menu Table", "Chọn Table → Insert Table rồi chọn kích thước trên lưới."),
        ("Chèn và nhập", "Chèn nhanh bằng /table", "Gõ /table và nhấn Enter để tạo bảng 3×3 có hàng tiêu đề."),
        ("Chèn và nhập", "Dán bảng HTML", "Sao chép bảng từ nguồn có HTML rồi dán bằng Ctrl/Cmd+V."),
        ("Chèn và nhập", "Dán dữ liệu dạng TSV", "Sao chép vùng ô có cột phân cách bằng Tab rồi dán để tạo bảng."),
        ("Chèn và nhập", "Giữ cấu trúc khi dán", "UEditor cố giữ gộp ô, độ rộng, chiều cao, màu, viền và định dạng chữ nếu nguồn cung cấp."),
        ("Chọn và điều hướng", "Bấm để nhập trong ô", "Bấm một ô rồi nhập hoặc sửa nội dung như văn bản thông thường."),
        ("Chọn và điều hướng", "Nhấp đúp chuột trái lên chữ", "Bôi đen phần chữ trong ô để định dạng nội dung."),
        ("Chọn và điều hướng", "Nhấp chuột trái ba lần lên chữ", "Chọn nguyên ô thay vì chỉ bôi đen phần chữ."),
        ("Chọn và điều hướng", "Nhấp đúp chuột trái vào vùng trống", "Chọn nguyên ô và mở thanh định dạng ô khi menu nổi được bật."),
        ("Chọn và điều hướng", "Chọn vùng nhiều ô", "Kéo qua các ô liền kề để tạo vùng chọn hình chữ nhật."),
        ("Chọn và điều hướng", "Sang ô kế tiếp", "Nhấn Tab."),
        ("Chọn và điều hướng", "Về ô trước", "Nhấn Shift+Tab."),
        ("Chọn và điều hướng", "Tạo hàng ở ô cuối", "Nhấn Tab tại ô cuối cùng của bảng."),
        ("Chọn và điều hướng", "Cuộn ngang bảng rộng", "Dùng thanh cuộn ngang bên dưới bảng khi bảng rộng hơn vùng soạn thảo."),
        ("Hàng và cột", "Thêm cột bên trái", "Dùng Add Column Before trên toolbar, menu bảng hoặc menu tay nắm cột."),
        ("Hàng và cột", "Thêm cột bên phải", "Dùng Add Column After trên toolbar, menu bảng hoặc menu tay nắm cột."),
        ("Hàng và cột", "Thêm hàng phía trên", "Dùng Add Row Before trên toolbar, menu bảng hoặc menu tay nắm hàng."),
        ("Hàng và cột", "Thêm hàng phía dưới", "Dùng Add Row After trên toolbar, menu bảng hoặc menu tay nắm hàng."),
        ("Hàng và cột", "Nhân đôi cột", "Mở menu tay nắm cột và chọn Duplicate Column."),
        ("Hàng và cột", "Nhân đôi hàng", "Mở menu tay nắm hàng và chọn Duplicate Row."),
        ("Hàng và cột", "Xóa dữ liệu cột", "Chọn Clear Column Contents; cấu trúc cột vẫn được giữ."),
        ("Hàng và cột", "Xóa dữ liệu hàng", "Chọn Clear Row Contents; cấu trúc hàng vẫn được giữ."),
        ("Hàng và cột", "Xóa cột", "Chọn Delete Column; cả dữ liệu và cấu trúc cột bị xóa."),
        ("Hàng và cột", "Xóa hàng", "Chọn Delete Row; cả dữ liệu và cấu trúc hàng bị xóa."),
        ("Hàng và cột", "Xóa toàn bộ bảng", "Mở menu ba chấm hoặc menu Table và chọn Delete Table."),
        ("Hàng và cột", "Đổi vị trí cột", "Kéo tay nắm phía trên cột sang vị trí mới."),
        ("Hàng và cột", "Đổi vị trí hàng", "Kéo tay nắm bên trái hàng sang vị trí mới."),
        ("Hàng và cột", "Thêm nhanh một cột", "Rê mép phải bảng rồi bấm thanh dấu +."),
        ("Hàng và cột", "Thêm nhanh nhiều cột", "Kéo thanh dấu + ở mép phải sang phải."),
        ("Hàng và cột", "Thêm nhanh một hàng", "Rê mép dưới bảng rồi bấm thanh dấu +."),
        ("Hàng và cột", "Thêm nhanh nhiều hàng", "Kéo thanh dấu + ở mép dưới xuống dưới."),
        ("Tiêu đề và vị trí", "Bật/tắt hàng tiêu đề", "Mở menu ba chấm và chọn Toggle Header Row."),
        ("Tiêu đề và vị trí", "Bật/tắt cột tiêu đề", "Mở menu ba chấm và chọn Toggle Header Column."),
        ("Tiêu đề và vị trí", "Căn bảng bên trái", "Chọn Align Table Left."),
        ("Tiêu đề và vị trí", "Căn bảng ở giữa", "Chọn Align Table Center."),
        ("Tiêu đề và vị trí", "Căn bảng bên phải", "Chọn Align Table Right."),
        ("Kích thước", "Đổi chiều rộng cột", "Rê mép phải ô đến khi hiện tay kéo rồi kéo ngang."),
        ("Kích thước", "Đổi chiều cao hàng", "Rê mép dưới hàng đến khi hiện đường hướng dẫn rồi kéo dọc."),
        ("Kích thước", "Đổi kích thước toàn bảng", "Kéo nút chéo ở góc dưới bên phải bảng."),
        ("Kích thước", "Khóa resize theo một trục", "Giữ Ctrl trong khi kéo nút đổi kích thước toàn bảng."),
        ("Kích thước", "Giữ tỷ lệ bảng", "Giữ Ctrl+Shift trong khi kéo nút đổi kích thước toàn bảng."),
        ("Kích thước", "Xem kích thước khi kéo", "Theo dõi nhãn rộng × cao xuất hiện trong lúc resize toàn bảng."),
        ("Ô và nội dung", "Gộp ô", "Chọn vùng ô liền kề hình chữ nhật rồi bấm Merge Cells."),
        ("Ô và nội dung", "Tách ô", "Chọn ô đã gộp rồi bấm Split Cell."),
        ("Ô và nội dung", "Màu nền ô", "Chọn một hoặc nhiều ô, bấm Cell Background và chọn màu."),
        ("Ô và nội dung", "Căn chữ trái", "Chọn ô/chữ rồi dùng Align Left trên thanh định dạng."),
        ("Ô và nội dung", "Căn chữ giữa", "Chọn ô/chữ rồi dùng Align Center."),
        ("Ô và nội dung", "Căn chữ phải", "Chọn ô/chữ rồi dùng Align Right."),
        ("Ô và nội dung", "Căn đều hai lề", "Chọn ô/chữ rồi dùng Justify."),
        ("Ô và nội dung", "Căn dọc phía trên", "Mở Vertical Alignment và chọn Top Align."),
        ("Ô và nội dung", "Căn dọc chính giữa", "Mở Vertical Alignment và chọn Middle Align."),
        ("Ô và nội dung", "Căn dọc phía dưới", "Mở Vertical Alignment và chọn Bottom Align."),
        ("Ô và nội dung", "Chữ nằm ngang", "Mở Text Direction và chọn Horizontal Text."),
        ("Ô và nội dung", "Chữ theo chiều dọc", "Mở Text Direction và chọn Vertical Text."),
        ("Ô và nội dung", "Cho phép xuống dòng", "Bật Wrap Text."),
        ("Ô và nội dung", "Không xuống dòng", "Chọn Don't Wrap Text; cột có thể rộng hơn nội dung hiển thị."),
        ("Viền ô", "Tất cả đường viền", "Trong Cell Border, chọn All Borders."),
        ("Viền ô", "Viền ngoài vùng chọn", "Chọn Outer Borders."),
        ("Viền ô", "Toàn bộ viền bên trong", "Chọn Inner Borders."),
        ("Viền ô", "Viền ngang bên trong", "Chọn Inner Horizontal Borders."),
        ("Viền ô", "Viền dọc bên trong", "Chọn Inner Vertical Borders."),
        ("Viền ô", "Viền trên", "Chọn Top Border."),
        ("Viền ô", "Viền phải", "Chọn Right Border."),
        ("Viền ô", "Viền dưới", "Chọn Bottom Border."),
        ("Viền ô", "Viền trái", "Chọn Left Border."),
        ("Viền ô", "Xóa viền", "Chọn Clear Border."),
        ("Viền ô", "Chế độ vẽ viền", "Chọn Draw Borders rồi chọn vị trí viền cần áp dụng."),
        ("Viền ô", "Chế độ ẩn viền", "Chọn Hide Borders rồi chọn vị trí cần bỏ."),
        ("Viền ô", "Kiểu liền", "Chọn Solid."),
        ("Viền ô", "Kiểu nét đứt", "Chọn Dashed."),
        ("Viền ô", "Kiểu chấm", "Chọn Dotted."),
        ("Viền ô", "Kiểu hai đường", "Chọn Double."),
        ("Viền ô", "Độ dày 1 px", "Chọn 1px."),
        ("Viền ô", "Độ dày 2 px", "Chọn 2px."),
        ("Viền ô", "Độ dày 3 px", "Chọn 3px."),
        ("Viền ô", "Độ dày 4 px", "Chọn 4px."),
        ("Viền ô", "Màu trong bảng màu", "Bấm Border Color và chọn mẫu màu có sẵn."),
        ("Viền ô", "Màu tùy chọn", "Trong bảng màu, chọn More Colors và chọn màu hệ thống."),
        ("Công thức", "Bắt đầu bằng dấu =", "Trong ô bảng, gõ = để mở gợi ý hàm hoặc nhập biểu thức."),
        ("Công thức", "Mở bảng công thức của ô", "Chọn ô rồi bấm Formula trên thanh định dạng ô."),
        ("Công thức", "Hiển thị thanh công thức", "Chọn một ô đã có công thức."),
        ("Công thức", "Sửa trong bảng công thức", "Nhập công thức vào ô Formula và bấm Apply."),
        ("Công thức", "Sửa trên thanh công thức", "Sửa biểu thức trong trường Edit Formula."),
        ("Công thức", "Tự thêm dấu =", "Khi nhập công thức trong bảng Formula mà thiếu =, UEditor tự bổ sung."),
        ("Công thức", "Áp dụng bằng nút Apply", "Bấm Apply trong bảng công thức."),
        ("Công thức", "Áp dụng bằng Enter trong ô", "Nhấn Enter khi đã nhập xong biểu thức trực tiếp."),
        ("Công thức", "Áp dụng bằng Enter trên thanh", "Nhấn Enter trong trường Edit Formula."),
        ("Công thức", "Hủy nhập trực tiếp", "Nhấn Escape khi đang gõ công thức trong ô."),
        ("Công thức", "Hủy sửa trên thanh", "Nhấn Escape trong trường Edit Formula để bỏ bản nháp."),
        ("Công thức", "Xóa công thức bằng Clear", "Bấm Clear trong bảng công thức."),
        ("Công thức", "Xóa ô bằng nút thùng rác", "Bấm Clear Cell trên thanh công thức."),
        ("Công thức", "Xóa bằng bàn phím", "Chọn ô công thức rồi nhấn Backspace hoặc Delete."),
        ("Công thức", "Chuyển thành giá trị", "Bấm Convert to Value (#); kết quả trở thành dữ liệu cố định."),
        ("Công thức", "Tính lại thủ công", "Bấm Recalculate trong bảng công thức."),
        ("Công thức", "Xem địa chỉ ô", "Địa chỉ như B2 xuất hiện ở đầu thanh công thức."),
        ("Công thức", "Tham chiếu một ô", "Nhập địa chỉ như A2."),
        ("Công thức", "Tham chiếu một vùng", "Nhập vùng như A2:C5."),
        ("Công thức", "Chọn ô tham chiếu bằng chuột", "Trong lúc gõ công thức, bấm một ô khác trong cùng bảng."),
        ("Công thức", "Chọn vùng bằng chuột", "Trong lúc gõ công thức, kéo qua các ô để chèn địa chỉ vùng."),
        ("Công thức", "Tự phân tách nhiều tham chiếu", "Khi cần, UEditor chèn dấu phẩy giữa các tham chiếu trong hàm."),
        ("Công thức", "Chặn tham chiếu gây vòng lặp", "Vùng gây vòng lặp được đánh dấu cảnh báo và không được chèn."),
        ("Công thức", "Phép cộng", "Dùng toán tử +."),
        ("Công thức", "Phép trừ", "Dùng toán tử -."),
        ("Công thức", "Phép nhân", "Dùng toán tử *."),
        ("Công thức", "Phép chia", "Dùng toán tử /."),
        ("Công thức", "Thứ tự tính bằng ngoặc", "Dùng cặp ngoặc tròn ()."),
        ("Công thức", "Số âm và số thập phân", "Nhập ví dụ -1.5 hoặc .25."),
        ("Công thức", "Hàm SUM", "Tính tổng một hoặc nhiều giá trị/vùng."),
        ("Công thức", "Hàm AVG", "Tính trung bình các giá trị số."),
        ("Công thức", "Hàm MIN", "Lấy giá trị nhỏ nhất."),
        ("Công thức", "Hàm MAX", "Lấy giá trị lớn nhất."),
        ("Công thức", "Hàm COUNT", "Đếm ô đọc được như số."),
        ("Công thức", "Mở danh sách hàm", "Gõ = trong ô bảng."),
        ("Công thức", "Lọc danh sách hàm", "Gõ tiếp chữ cái đầu của tên hàm."),
        ("Công thức", "Di chuyển trong danh sách hàm", "Nhấn Arrow Up hoặc Arrow Down."),
        ("Công thức", "Chèn hàm từ gợi ý", "Nhấn Tab hoặc bấm tên hàm."),
        ("Công thức", "Tự tính lại khi dữ liệu đổi", "UEditor tính lại công thức phụ thuộc sau khi ô nguồn thay đổi."),
        ("Công thức", "Tính lại khi mở nội dung", "Các bảng được tính lại khi nội dung đã lưu được nạp vào editor."),
        ("Công thức", "Lỗi #EMPTY", "Báo công thức trống."),
        ("Công thức", "Lỗi #INVALID-REFERENCE", "Báo tham chiếu không hợp lệ hoặc không đọc được như số."),
        ("Công thức", "Lỗi #INVALID-FORMULA", "Báo sai cú pháp hoặc dùng hàm không hỗ trợ."),
        ("Công thức", "Lỗi #DIVISION-BY-ZERO", "Báo phép chia cho 0 hoặc AVG không có số."),
        ("Công thức", "Lỗi #CIRCULAR-REFERENCE", "Báo vòng tham chiếu trực tiếp hoặc gián tiếp."),
        ("Công thức", "Định dạng Text", "Hiển thị giá trị thô."),
        ("Công thức", "Định dạng Number", "Hiển thị số có phân cách hàng nghìn."),
        ("Công thức", "Định dạng Currency", "Hiển thị tiền tệ USD."),
        ("Công thức", "Định dạng Percent", "Hiển thị giá trị theo phần trăm."),
        ("Công thức", "Định dạng Date", "Hiển thị ngày từ số serial."),
        ("Công thức", "Chỉ báo lỗi trên thanh", "Thanh công thức hiện Formula Error khi ô đang có lỗi."),
    ]
    assert len(items) == 131, f"Expected 131 Table entries, found {len(items)}"
    return [[f"T{index:03d}", group, feature, usage] for index, (group, feature, usage) in enumerate(items, start=1)]


def get_core_feature_rows() -> list[list[str]]:
    """Non-table user-visible operations/options for the final coverage index."""
    items = [
        ("Soạn thảo", "Nhập rich text", "Bấm vùng nội dung và nhập văn bản."),
        ("Soạn thảo", "Chọn nội dung", "Kéo chuột hoặc dùng bàn phím để bôi đen."),
        ("Soạn thảo", "Xóa nội dung", "Dùng Backspace hoặc Delete."),
        ("Soạn thảo", "Cắt", "Ctrl/Cmd+X hoặc Edit → Cut."),
        ("Soạn thảo", "Sao chép", "Ctrl/Cmd+C hoặc Edit → Copy."),
        ("Soạn thảo", "Dán có định dạng", "Ctrl/Cmd+V; kết quả phụ thuộc dữ liệu clipboard."),
        ("Soạn thảo", "Dán dạng văn bản", "Edit → Paste as Text."),
        ("Soạn thảo", "Chọn toàn bộ", "Ctrl/Cmd+A hoặc Edit → Select All."),
        ("Soạn thảo", "Hoàn tác", "Ctrl/Cmd+Z hoặc Undo."),
        ("Soạn thảo", "Làm lại", "Ctrl/Cmd+Y hoặc Shift+Cmd+Z."),
        ("Soạn thảo", "Đếm từ", "Xem bộ đếm dưới editor khi màn hình bật tính năng."),
        ("Soạn thảo", "Đếm ký tự", "Xem bộ đếm dưới editor khi màn hình bật tính năng."),
        ("Soạn thảo", "Giới hạn ký tự", "Theo dõi bộ đếm/giới hạn do màn hình cấu hình."),
        ("Soạn thảo", "Chế độ chỉ đọc", "Nội dung hiển thị nhưng không cho chỉnh sửa."),
        ("Soạn thảo", "Placeholder", "Thông báo gợi ý xuất hiện khi editor trống."),
        ("Kiểu chữ", "Chọn font", "Mở Font Family và chọn font được cấu hình."),
        ("Kiểu chữ", "Trả về font mặc định", "Chọn Default trong Font Family."),
        ("Kiểu chữ", "Chọn cỡ chữ có sẵn", "Mở Font Size và chọn giá trị."),
        ("Kiểu chữ", "Nhập cỡ chữ 8–96 px", "Mở Text Style, nhập số rồi xác nhận."),
        ("Kiểu chữ", "Normal", "Chuyển khối hiện tại về đoạn văn thường."),
        ("Kiểu chữ", "Heading 1", "Chọn Heading 1 hoặc dùng Alt+1."),
        ("Kiểu chữ", "Heading 2", "Chọn Heading 2 hoặc dùng Alt+2."),
        ("Kiểu chữ", "Heading 3", "Chọn Heading 3 hoặc dùng Alt+3."),
        ("Kiểu chữ", "Line Height mặc định", "Mở Line Height và chọn Default."),
        ("Kiểu chữ", "Chọn giãn dòng", "Mở Line Height và chọn mức được cấu hình."),
        ("Kiểu chữ", "Letter Spacing mặc định", "Mở Letter Spacing và chọn Default."),
        ("Kiểu chữ", "Chọn giãn chữ", "Mở Letter Spacing và chọn mức được cấu hình."),
        ("Định dạng chữ", "In đậm", "Ctrl/Cmd+B hoặc Bold."),
        ("Định dạng chữ", "In nghiêng", "Ctrl/Cmd+I hoặc Italic."),
        ("Định dạng chữ", "Gạch chân", "Ctrl/Cmd+U hoặc Underline."),
        ("Định dạng chữ", "Gạch ngang", "Ctrl/Cmd+Shift+S hoặc Strikethrough."),
        ("Định dạng chữ", "Mã nội dòng", "Ctrl/Cmd+E hoặc Inline Code."),
        ("Định dạng chữ", "Chỉ số dưới", "Ctrl/Cmd+, hoặc Subscript."),
        ("Định dạng chữ", "Chỉ số trên", "Ctrl/Cmd+. hoặc Superscript."),
        ("Màu", "Màu chữ có sẵn", "Mở Text Color và chọn mẫu màu."),
        ("Màu", "Màu chữ tùy chọn", "Chọn More Colors trong bảng màu chữ."),
        ("Màu", "Màu đánh dấu có sẵn", "Mở Highlight và chọn mẫu màu."),
        ("Màu", "Màu đánh dấu tùy chọn", "Chọn More Colors trong bảng Highlight."),
        ("Định dạng chữ", "Xóa định dạng", "Format → Clear Formatting."),
        ("Đoạn văn", "Căn trái", "Mở Alignment và chọn Align Left."),
        ("Đoạn văn", "Căn giữa", "Mở Alignment và chọn Align Center."),
        ("Đoạn văn", "Căn phải", "Mở Alignment và chọn Align Right."),
        ("Đoạn văn", "Căn đều", "Mở Alignment và chọn Justify."),
        ("Đoạn văn", "Tăng lề", "Bấm Increase Indent hoặc Tab trong ngữ cảnh phù hợp."),
        ("Đoạn văn", "Giảm lề", "Bấm Decrease Indent hoặc Shift+Tab."),
        ("Danh sách", "Danh sách dấu đầu dòng", "Bấm Bullet List hoặc Ctrl/Cmd+Shift+8."),
        ("Danh sách", "Danh sách đánh số", "Chọn Numbered List hoặc Ctrl/Cmd+Shift+7."),
        ("Danh sách", "Danh sách công việc", "Chọn Task List hoặc Ctrl/Cmd+Shift+9."),
        ("Danh sách", "Tăng cấp danh sách", "Nhấn Tab tại một mục danh sách."),
        ("Danh sách", "Giảm cấp danh sách", "Nhấn Shift+Tab tại một mục danh sách."),
        ("Khối nội dung", "Trích dẫn", "Chọn Quote hoặc Ctrl/Cmd+Shift+B."),
        ("Khối nội dung", "Khối mã", "Chọn Code Block hoặc Ctrl/Cmd+Alt+C."),
        ("Khối nội dung", "Chọn ngôn ngữ mã", "Mở danh sách ngôn ngữ trên Code Block."),
        ("Khối nội dung", "Sao chép khối mã", "Rê lên Code Block và bấm Copy."),
        ("Khối nội dung", "Đường phân cách", "Chọn Divider, Horizontal Rule hoặc gõ ---."),
        ("Liên kết", "Tạo liên kết", "Bôi đen chữ, bấm Link và nhập URL."),
        ("Liên kết", "Xem trước liên kết", "Bấm hoặc đặt con trỏ tại liên kết trong chế độ sửa."),
        ("Liên kết", "Mở liên kết", "Chọn Open Link trong menu nổi."),
        ("Liên kết", "Sửa liên kết", "Chọn Edit Link và nhập URL mới."),
        ("Liên kết", "Gỡ liên kết", "Chọn Remove Link."),
        ("Liên kết", "Kiểm tra URL an toàn", "UEditor từ chối giao thức/URL không hợp lệ."),
        ("Liên kết", "Bookmark Card", "Gõ /bookmark, nhập URL và xác nhận."),
        ("Hình ảnh", "Chèn ảnh từ URL", "Mở Insert Image → URL, nhập địa chỉ và Alt text."),
        ("Hình ảnh", "Tải ảnh từ thiết bị", "Mở Insert Image → Upload và chọn tệp."),
        ("Hình ảnh", "Dán ảnh", "Sao chép ảnh rồi nhấn Ctrl/Cmd+V."),
        ("Hình ảnh", "Thả ảnh", "Kéo tệp ảnh vào vị trí cần chèn."),
        ("Hình ảnh", "Alt text", "Nhập mô tả trong hộp chèn ảnh."),
        ("Hình ảnh", "Kéo đổi kích thước", "Chọn ảnh và kéo tay nắm."),
        ("Hình ảnh", "Bố cục Block", "Chọn ảnh → Image Layout Block."),
        ("Hình ảnh", "Bọc chữ bên trái", "Chọn Image Layout Left."),
        ("Hình ảnh", "Bọc chữ bên phải", "Chọn Image Layout Right."),
        ("Hình ảnh", "Cỡ S", "Chọn Image Width Small."),
        ("Hình ảnh", "Cỡ M", "Chọn Image Width Medium."),
        ("Hình ảnh", "Cỡ L", "Chọn Image Width Large."),
        ("Hình ảnh", "Khôi phục kích thước", "Chọn Reset Image Size."),
        ("Hình ảnh", "Xóa ảnh", "Chọn Delete Image."),
        ("Tệp", "Đính kèm tệp", "Gõ /file và chọn tệp."),
        ("Tệp", "Tải lại hoặc tải xuống", "Dùng Retry Upload khi lỗi; bấm thẻ tệp để tải xuống."),
        ("Lệnh nhanh", "Mở menu /", "Gõ / ở đầu dòng trống."),
        ("Lệnh nhanh", "Lọc lệnh", "Gõ tiếp tên khối cần chèn."),
        ("Lệnh nhanh", "Điều khiển bằng bàn phím", "Arrow Up/Down để chọn, Enter để chèn, Esc để đóng."),
        ("Lệnh nhanh", "Floating Add Block", "Bấm Add block tại dòng trống khi nút này xuất hiện."),
        ("Callout", "Chèn Callout", "Gõ /callout hoặc chọn lệnh tương ứng."),
        ("Callout", "Đổi emoji", "Bấm emoji ở đầu Callout và chọn biểu tượng."),
        ("Callout", "Đổi màu nền", "Mở Background Color của Callout."),
        ("Biểu mẫu", "Checkbox vuông", "Chọn Form Checkbox hoặc /checkbox."),
        ("Biểu mẫu", "Checkbox tròn", "Chọn Round Checkbox hoặc /round checkbox."),
        ("Biểu mẫu", "Radio và Group Name", "Chèn Form Radio, đặt cùng Group Name cho các lựa chọn cùng nhóm."),
        ("Emoji", "Mở gợi ý bằng dấu :", "Gõ : kèm từ khóa trong nội dung."),
        ("Emoji", "Tìm emoji", "Nhập từ khóa trong bảng emoji."),
        ("Emoji", "Di chuyển trong lưới", "Dùng các phím mũi tên."),
        ("Emoji", "Chèn emoji", "Nhấn Enter hoặc bấm emoji."),
        ("Emoji", "Đóng bảng emoji", "Nhấn Escape."),
        ("Menu và đầu ra", "Menu cổ điển", "Sử dụng File, Edit, View, Insert, Format, Tools và Table khi được bật."),
        ("Menu và đầu ra", "Preview", "Chọn View → Preview hoặc bấm biểu tượng con mắt."),
        ("Menu và đầu ra", "Source Code", "Chọn View/Tools → Source Code."),
        ("Menu và đầu ra", "Áp dụng HTML", "Trong Source Code, bấm Apply sau khi sửa."),
        ("Menu và đầu ra", "Fullscreen", "Chọn View → Fullscreen; nhấn Esc để thoát."),
        ("Menu và đầu ra", "Export HTML", "Chọn File → Export HTML."),
        ("Menu và đầu ra", "Save", "Chọn File → Save khi màn hình đã tích hợp hành động lưu."),
    ]
    assert len(items) == 100, f"Expected 100 core entries, found {len(items)}"
    return [[f"U{index:03d}", group, feature, usage] for index, (group, feature, usage) in enumerate(items, start=1)]


def add_cover(document: Document) -> None:
    document.add_paragraph()
    band = document.add_table(rows=1, cols=1)
    band.alignment = WD_TABLE_ALIGNMENT.CENTER
    band.autofit = False
    band.cell(0, 0).width = Cm(17.2)
    set_cell_shading(band.cell(0, 0), BLUE)
    set_cell_margins(band.cell(0, 0), top=110, bottom=110, start=160, end=160)
    paragraph = band.cell(0, 0).paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("UNDERVERSE  •  TÀI LIỆU HƯỚNG DẪN")
    set_run_typeface(run)
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(WHITE)

    document.add_paragraph().paragraph_format.space_after = Pt(22)
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("HƯỚNG DẪN SỬ DỤNG\nUEDITOR")

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.add_run("Tài liệu hướng dẫn chi tiết dành cho người dùng cuối")

    line = document.add_table(rows=1, cols=2)
    line.alignment = WD_TABLE_ALIGNMENT.LEFT
    line.autofit = False
    set_column_widths(line, [1.0, 16.0])
    set_cell_shading(line.cell(0, 0), BLUE)
    set_cell_shading(line.cell(0, 1), BLUE_LIGHT)
    for cell in line.row_cells(0):
        set_cell_margins(cell, top=22, bottom=22, start=0, end=0)

    document.add_paragraph().paragraph_format.space_after = Pt(18)
    intro = document.add_paragraph()
    intro.paragraph_format.space_after = Pt(8)
    intro.add_run(
        "Soạn thảo văn bản • Định dạng • Hình ảnh • Liên kết • Khối nội dung • "
        "Bảng nâng cao • Công thức • Phím tắt • Xử lý lỗi"
    ).bold = True

    add_callout(
        document,
        "info",
        "Ngôn ngữ của tài liệu",
        "Phần giải thích được viết bằng tiếng Việt. Theo yêu cầu, toàn bộ ảnh chụp minh họa dùng giao diện tiếng Anh. "
        "Tên tiếng Anh quan trọng được đặt trong ngoặc để người đọc dễ đối chiếu.",
    )

    document.add_paragraph().paragraph_format.space_after = Pt(30)
    meta = document.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.autofit = False
    set_column_widths(meta, [5.2, 11.5])
    values = [
        ("Phiên bản tài liệu", "2.0"),
        ("Ngày cập nhật", "24/07/2026"),
        ("Đối tượng", "Khách hàng / người dùng cuối"),
        ("Phạm vi", "UEditor trong dự án Underverse"),
    ]
    for idx, (label, value) in enumerate(values):
        left, right = meta.rows[idx].cells
        left.text = label
        right.text = value
        set_cell_shading(left, SLATE_LIGHT)
        for cell in (left, right):
            set_cell_margins(cell, top=90, start=110, bottom=90, end=110)
            set_cell_border(cell, bottom={"val": "single", "sz": "4", "color": GRAY_BORDER})
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)
        left.paragraphs[0].runs[0].font.bold = True
        left.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(SLATE)

    document.add_paragraph()
    end = document.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    end_run = end.add_run("HƯỚNG DẪN SỬ DỤNG  •  PHIÊN BẢN 2.0")
    end_run.font.size = Pt(8.5)
    end_run.font.bold = True
    end_run.font.color.rgb = RGBColor.from_string(BLUE)


def build_document() -> Document:
    document = configure_document()
    add_cover(document)

    add_heading(document, "Thông tin và phạm vi tài liệu", 1, page_break=True)
    add_paragraph(
        document,
        "Tài liệu này hướng dẫn cách sử dụng UEditor ở góc nhìn người dùng cuối. "
        "Nội dung được đối chiếu trực tiếp với phần triển khai UEditor trong dự án, bao gồm thanh công cụ, "
        "menu cổ điển, lệnh nhanh, bảng, công thức, hình ảnh, tệp đính kèm và các phím tắt."
    )
    add_callout(
        document,
        "warning",
        "Giao diện có thể khác đôi chút",
        "UEditor cho phép hệ thống bật/tắt thanh công cụ, menu cổ điển, menu nổi, bộ đếm và một số chức năng tích hợp. "
        "Vì vậy, một màn hình nghiệp vụ có thể không hiển thị toàn bộ nút được mô tả trong tài liệu.",
    )
    add_heading(document, "Cách đọc tên chức năng", 2)
    add_bullets(
        document,
        [
            "Tên tiếng Việt mô tả ý nghĩa sử dụng; tên tiếng Anh trong ngoặc là nhãn xuất hiện trên ảnh minh họa.",
            "Ctrl/Cmd nghĩa là dùng Ctrl trên Windows/Linux và dùng ⌘ Command trên macOS.",
            "“Chọn nội dung” nghĩa là bôi đen chữ, chọn ảnh hoặc chọn một hay nhiều ô bảng trước khi dùng lệnh.",
            "Nút màu xám hoặc không bấm được thường có nghĩa là vị trí con trỏ hiện tại chưa phù hợp với lệnh.",
        ],
    )
    add_heading(document, "Quy ước mức độ đảm bảo của phím tắt", 2)
    add_table(
        document,
        ["Loại", "Ý nghĩa"],
        [
            ["Phím tắt UEditor", "Được UEditor/Tiptap xử lý trực tiếp và dùng được khi con trỏ đang ở vùng soạn thảo."],
            ["Phím tắt trình duyệt", "Cắt, chép, dán… do trình duyệt/hệ điều hành xử lý; có thể phụ thuộc quyền clipboard."],
            ["Nhãn phím tắt trong menu", "Menu có thể hiển thị gợi ý như Ctrl+S, Ctrl+K, Ctrl+Shift+F; khả năng hoạt động còn phụ thuộc màn hình tích hợp."],
        ],
        [4.0, 12.7],
    )

    add_heading(document, "Mục lục", 1, page_break=True)
    toc_paragraph = document.add_paragraph()
    add_toc(toc_paragraph)
    add_callout(
        document,
        "info",
        "Cập nhật mục lục",
        "Nếu số trang chưa hiện ngay, mở tài liệu bằng Microsoft Word, bấm chuột phải vào mục lục và chọn Update Field → Update entire table.",
    )

    add_heading(document, "1. Bắt đầu nhanh trong 3 phút", 1, page_break=True)
    add_heading(document, "1.1 Quy trình soạn và lưu nội dung", 2)
    add_numbered_steps(
        document,
        [
            "Bấm vào vùng trắng của UEditor để đặt con trỏ.",
            "Nhập văn bản như trong trình soạn thảo thông thường; nhấn Enter để tạo đoạn mới.",
            "Bôi đen phần cần định dạng, sau đó dùng thanh công cụ hoặc menu nổi.",
            "Gõ “/” ở một đoạn để chèn nhanh tiêu đề, danh sách, bảng, hộp thông tin, tệp…",
            "Kiểm tra bộ đếm từ/ký tự ở chân trình soạn thảo nếu màn hình có bật giới hạn.",
            "Bấm nút Lưu của màn hình nghiệp vụ hoặc File → Save nếu chức năng này đã được hệ thống tích hợp.",
            "Chờ thông báo lưu thành công, đặc biệt khi nội dung có ảnh hoặc tệp cần tải lên.",
        ],
    )
    add_callout(
        document,
        "tip",
        "Cách thao tác nhanh nhất",
        "Với nội dung thông thường, chỉ cần nhớ 4 thao tác: bôi đen để định dạng, gõ “/” để chèn khối, "
        "gõ “:” để tìm emoji và dùng Ctrl/Cmd+Z để hoàn tác.",
    )
    add_heading(document, "1.2 Bảng tra cứu nhanh", 2)
    add_table(
        document,
        ["Nhu cầu", "Cách làm nhanh"],
        [
            ["Tạo tiêu đề", "Gõ /heading rồi chọn cấp 1–3; hoặc Ctrl/Cmd+Alt+1, 2, 3."],
            ["In đậm / nghiêng / gạch chân", "Bôi đen chữ rồi nhấn Ctrl/Cmd+B, I hoặc U."],
            ["Chèn liên kết", "Bôi đen chữ → nút Link → nhập URL hợp lệ → xác nhận."],
            ["Chèn ảnh", "Nút Image → URL/Upload; cũng có thể dán hoặc kéo thả ảnh vào editor."],
            ["Chèn bảng", "Nút Table → rê chọn số ô; hoặc gõ /table để tạo bảng 3×3."],
            ["Tính tổng trong bảng", "Chọn ô kết quả → nhập =SUM(A1:A3) → Enter."],
            ["Chèn hộp thông tin", "Gõ /callout → Enter → chọn emoji và màu nền."],
            ["Chèn tệp", "Gõ /file → Enter → chọn tệp trên thiết bị."],
            ["Xem trước", "View → Preview hoặc nút hình con mắt."],
            ["Quay lại thao tác trước", "Ctrl/Cmd+Z."],
        ],
        [5.0, 11.7],
    )

    add_heading(document, "2. Làm quen với giao diện", 1, page_break=True)
    add_heading(document, "2.1 Các vùng chính", 2)
    add_table(
        document,
        ["Vùng", "Mục đích", "Khi nào xuất hiện"],
        [
            ["Menu cổ điển (Menu bar)", "File, Edit, View, Insert, Format, Tools, Table.", "Chỉ xuất hiện khi màn hình bật chế độ menu cổ điển."],
            ["Thanh công cụ (Toolbar)", "Định dạng chữ, màu, căn lề, danh sách, ảnh, bảng, hoàn tác…", "Thường nằm trên vùng soạn thảo; có thể cuộn ngang ở màn hình hẹp."],
            ["Vùng soạn thảo", "Nhập, chọn, dán và sắp xếp nội dung.", "Luôn có khi editor ở chế độ chỉnh sửa hoặc xem."],
            ["Menu nổi (Bubble menu)", "Công cụ theo đối tượng đang chọn: chữ, ảnh, liên kết hoặc ô bảng.", "Xuất hiện sau khi bôi đen/chọn đối tượng và editor đang có tiêu điểm."],
            ["Thanh công thức", "Hiển thị địa chỉ ô, công thức và các nút áp dụng/chuyển thành giá trị/xóa.", "Xuất hiện khi chọn ô bảng đang chứa công thức."],
            ["Điều khiển bảng", "Tay nắm hàng/cột, nút ba chấm, thanh thêm hàng/cột, tay nắm đổi kích thước.", "Hiện khi trỏ chuột hoặc đặt con trỏ trong bảng."],
            ["Chân editor", "Hiển thị số từ, số ký tự và giới hạn nếu có.", "Khi màn hình bật footer và bộ đếm."],
        ],
        [4.0, 7.3, 5.4],
        compact=True,
    )
    add_image(
        document,
        "01-menu-co-dien.png",
        "Hình 1. Menu cổ điển, thanh công cụ, vùng nhập và bộ đếm (giao diện tiếng Anh).",
        16.5,
    )
    add_image(
        document,
        "02-thanh-cong-cu.png",
        "Hình 2. Thanh công cụ đầy đủ: font, cỡ chữ, kiểu chữ, liên kết, màu và căn lề (giao diện tiếng Anh).",
        16.5,
    )
    add_heading(document, "2.2 Trạng thái nút", 2)
    add_bullets(
        document,
        [
            "Nút đang được tô nền: định dạng hiện đang áp dụng cho vị trí con trỏ hoặc vùng chọn.",
            "Nút mờ/khóa: thao tác không hợp lệ trong ngữ cảnh hiện tại; ví dụ Merge Cells chỉ dùng khi đã chọn nhiều ô liền kề.",
            "Nút có mũi tên hoặc bấm mở bảng chọn: chứa nhiều tùy chọn cùng nhóm.",
            "Biểu tượng có thể thay đổi theo đối tượng; ví dụ nút Merge Cells chuyển thành Split Cell khi đang ở ô đã gộp.",
        ],
    )
    add_heading(document, "2.3 Chế độ chỉ đọc", 2)
    add_paragraph(
        document,
        "Khi UEditor ở chế độ chỉ đọc, thanh công cụ, viền chỉnh sửa và tay nắm đổi kích thước thường bị ẩn. "
        "Người dùng vẫn có thể đọc nội dung, mở liên kết hoặc tải tệp nếu màn hình cho phép, nhưng không thể sửa."
    )

    add_heading(document, "3. Soạn thảo cơ bản", 1, page_break=True)
    add_heading(document, "3.1 Nhập, chọn và xóa nội dung", 2)
    add_table(
        document,
        ["Thao tác", "Cách thực hiện"],
        [
            ["Đặt con trỏ", "Bấm vào vị trí muốn nhập."],
            ["Tạo đoạn mới", "Nhấn Enter."],
            ["Chọn một từ", "Nhấp đúp vào từ."],
            ["Chọn một vùng chữ", "Giữ chuột trái và kéo qua nội dung; hoặc giữ Shift kết hợp phím mũi tên."],
            ["Chọn toàn bộ nội dung", "Ctrl/Cmd+A khi con trỏ đang ở trong editor."],
            ["Xóa vùng chọn", "Nhấn Backspace hoặc Delete."],
            ["Hoàn tác", "Ctrl/Cmd+Z hoặc nút Undo."],
            ["Làm lại", "Ctrl/Cmd+Y; trên macOS có thể dùng Shift+Cmd+Z."],
        ],
        [4.2, 12.5],
    )
    add_heading(document, "3.2 Cắt, sao chép và dán", 2)
    add_bullets(
        document,
        [
            "Cắt: Ctrl/Cmd+X.",
            "Sao chép: Ctrl/Cmd+C.",
            "Dán: Ctrl/Cmd+V.",
            "Dán như văn bản thuần: Edit → Paste as Text. Trình duyệt có thể hỏi quyền đọc clipboard.",
            "Dán bảng từ Excel/Google Sheets hoặc dữ liệu phân tách bằng tab: UEditor cố gắng chuyển thành bảng và giữ các dữ liệu/định dạng tương thích.",
            "Dán ảnh từ clipboard: ảnh được chèn tại vị trí con trỏ nếu định dạng và dung lượng được hệ thống chấp nhận.",
        ],
    )
    add_callout(
        document,
        "warning",
        "Quyền clipboard của trình duyệt",
        "Một số trình duyệt chặn nút Paste trong menu. Khi đó hãy dùng Ctrl/Cmd+V. "
        "Nếu dán bảng từ nguồn bên ngoài, luôn kiểm tra lại đường viền, gộp ô và công thức sau khi dán.",
    )
    add_heading(document, "3.3 Thụt lề bằng Tab", 2)
    add_bullets(
        document,
        [
            "Khi con trỏ ở giữa đoạn văn: Tab chèn một khoảng tab tại vị trí con trỏ; Shift+Tab xóa tab ngay trước con trỏ.",
            "Khi con trỏ ở đầu đoạn hoặc đang chọn nhiều đoạn: Tab tăng thụt lề cả khối; Shift+Tab giảm thụt lề.",
            "Trong danh sách: Tab đưa mục thành cấp con; Shift+Tab đưa mục lên một cấp.",
            "Trong bảng: Tab chuyển sang ô kế tiếp; Shift+Tab quay lại ô trước.",
            "Backspace ở đầu một đoạn đang thụt lề sẽ giảm một cấp trước khi gộp đoạn.",
        ],
    )

    add_heading(document, "4. Định dạng văn bản và đoạn", 1, page_break=True)
    add_heading(document, "4.1 Chọn kiểu đoạn", 2)
    add_table(
        document,
        ["Kiểu", "Mục đích", "Cách dùng"],
        [
            ["Văn bản thường (Normal text)", "Nội dung thân bài.", "Đặt con trỏ trong đoạn → Text Style → Normal text."],
            ["Tiêu đề 1 (Heading 1)", "Tiêu đề lớn nhất của nội dung.", "Text Style → Heading 1 hoặc Ctrl/Cmd+Alt+1."],
            ["Tiêu đề 2 (Heading 2)", "Nhóm nội dung cấp 2.", "Text Style → Heading 2 hoặc Ctrl/Cmd+Alt+2."],
            ["Tiêu đề 3 (Heading 3)", "Nhóm nội dung cấp 3.", "Text Style → Heading 3 hoặc Ctrl/Cmd+Alt+3."],
            ["Trích dẫn (Quote/Blockquote)", "Làm nổi bật lời dẫn hoặc ghi chú trích dẫn.", "Nút Quote, /quote hoặc Ctrl/Cmd+Shift+B."],
            ["Khối mã (Code Block)", "Hiển thị đoạn mã có tô màu cú pháp.", "Nút Code Block, /code hoặc Ctrl/Cmd+Alt+C."],
        ],
        [4.5, 5.6, 6.6],
    )
    add_callout(
        document,
        "tip",
        "Dùng tiêu đề đúng cấp",
        "Không chọn Heading chỉ để làm chữ to. Hãy dùng Heading 1 → Heading 2 → Heading 3 theo cấu trúc tài liệu để nội dung dễ đọc và hỗ trợ truy cập tốt hơn.",
    )
    add_heading(document, "4.2 Font, cỡ chữ, giãn dòng và giãn chữ", 2)
    add_bullets(
        document,
        [
            "Font Family: chọn font cho vùng chữ đã bôi đen hoặc chữ sẽ nhập tiếp theo. Danh sách font có thể được cấu hình theo từng màn hình.",
            "Font Size: các cỡ mặc định từ 8 đến 28, thêm 36, 48, 72 và 96 px. Menu nổi cho phép nhập trực tiếp giá trị 8–96 px.",
            "Line Height: giãn dòng mặc định hoặc các mức 1.2, 1.5, 1.75 và 2.",
            "Letter Spacing: giãn chữ từ -0.02em đến 0.08em; dùng giá trị âm để chữ sát hơn và dương để chữ thoáng hơn.",
            "Để trả về mặc định, chọn mục Default trong bảng chọn tương ứng.",
        ],
    )
    add_heading(document, "4.3 Định dạng ký tự", 2)
    add_table(
        document,
        ["Tính năng", "English label", "Phím tắt UEditor", "Cách dùng"],
        [
            ["In đậm", "Bold", "Ctrl/Cmd+B", "Bôi đen chữ rồi bấm B."],
            ["In nghiêng", "Italic", "Ctrl/Cmd+I", "Bôi đen chữ rồi bấm I."],
            ["Gạch chân", "Underline", "Ctrl/Cmd+U", "Bôi đen chữ rồi bấm U."],
            ["Gạch ngang", "Strikethrough", "Ctrl/Cmd+Shift+S", "Dùng cho nội dung đã hủy/không còn áp dụng."],
            ["Mã nội dòng", "Inline Code", "Ctrl/Cmd+E", "Dùng cho tên lệnh, biến hoặc đoạn mã ngắn."],
            ["Chỉ số dưới", "Subscript", "Ctrl/Cmd+,", "Ví dụ H₂O."],
            ["Chỉ số trên", "Superscript", "Ctrl/Cmd+.", "Ví dụ x²."],
        ],
        [3.4, 3.2, 4.2, 5.9],
        compact=True,
    )
    add_heading(document, "4.4 Màu chữ và màu đánh dấu", 2)
    add_numbered_steps(
        document,
        [
            "Bôi đen nội dung cần đổi màu.",
            "Bấm Text Color để chọn màu chữ hoặc Highlight để chọn màu nền chữ.",
            "Chọn một màu có sẵn; dùng More Colors nếu cần màu tùy chỉnh.",
            "Chọn Automatic/Default để trả về màu mặc định.",
        ],
    )
    add_callout(
        document,
        "warning",
        "Độ tương phản",
        "Không dùng màu quá nhạt cho chữ và không chỉ dùng màu để truyền đạt trạng thái. "
        "Khi gửi nội dung cho khách hàng, nên kiểm tra cả chế độ sáng và tối nếu hệ thống hỗ trợ.",
    )
    add_heading(document, "4.5 Căn lề, hướng chữ và xuống dòng", 2)
    add_bullets(
        document,
        [
            "Align Left / Center / Right / Justify: căn trái, giữa, phải hoặc đều hai bên cho đoạn/tiêu đề.",
            "Decrease/Increase Indent: giảm hoặc tăng thụt lề của đoạn.",
            "Khi con trỏ ở trong ô bảng, menu Alignment có thêm căn dọc Top/Middle/Bottom.",
            "Text Direction trong ô bảng: chữ ngang hoặc chữ dọc.",
            "Wrap Text: bật/tắt tự xuống dòng trong ô bảng.",
        ],
    )
    add_heading(document, "4.6 Xóa định dạng", 2)
    add_paragraph(
        document,
        "Trong menu cổ điển, chọn Format → Clear Formatting để bỏ các định dạng đang áp dụng và đưa khối về dạng cơ bản. "
        "Nếu chỉ muốn tắt một thuộc tính, bôi đen chữ rồi bấm lại nút đang được kích hoạt."
    )

    add_heading(document, "5. Danh sách, công việc và trường tương tác", 1, page_break=True)
    add_heading(document, "5.1 Danh sách dấu chấm và danh sách số", 2)
    add_bullets(
        document,
        [
            "Bullet List: dùng cho các ý không cần thứ tự; phím tắt Ctrl/Cmd+Shift+8.",
            "Numbered List: dùng cho quy trình hoặc thứ tự ưu tiên; phím tắt Ctrl/Cmd+Shift+7.",
            "Nhấn Enter để tạo mục tiếp theo.",
            "Nhấn Enter tại một mục trống để thoát danh sách.",
            "Tab tạo cấp con; Shift+Tab đưa mục lên cấp cha.",
        ],
    )
    add_heading(document, "5.2 Danh sách công việc (Task List)", 2)
    add_numbered_steps(
        document,
        [
            "Chọn nút danh sách → Task List, gõ /todo hoặc nhấn Ctrl/Cmd+Shift+9.",
            "Nhập nội dung công việc.",
            "Bấm vào ô vuông đầu dòng để đánh dấu hoàn thành/chưa hoàn thành.",
            "Dùng Tab/Shift+Tab để tạo công việc con hoặc đưa về cấp trên.",
        ],
    )
    add_heading(document, "5.3 Checkbox biểu mẫu", 2)
    add_paragraph(
        document,
        "Form Checkbox là một trường tương tác nằm ngay trong dòng chữ, khác với Task List. "
        "Có hai kiểu: ô vuông và ô tròn. Chọn từ thanh công cụ, Insert → Form Fields hoặc lệnh /checkbox và /round checkbox."
    )
    add_bullets(
        document,
        [
            "Bấm chuột vào checkbox để bật/tắt.",
            "Khi checkbox được focus bằng bàn phím, nhấn Space hoặc Enter để bật/tắt.",
            "Checkbox biểu mẫu hoạt động độc lập; nhiều checkbox có thể được chọn cùng lúc.",
        ],
    )
    add_heading(document, "5.4 Radio biểu mẫu", 2)
    add_numbered_steps(
        document,
        [
            "Gõ /radio hoặc chọn Insert → Form Fields → Form Radio Button.",
            "Bấm radio để chọn; các radio cùng Group Name chỉ cho phép chọn một giá trị.",
            "Chọn radio trong editor để hiện biểu tượng bánh răng.",
            "Bấm bánh răng và sửa Group Name nếu muốn nhiều radio thuộc cùng một nhóm.",
        ],
    )
    add_callout(
        document,
        "tip",
        "Đặt cùng tên nhóm",
        "Muốn tạo câu hỏi một lựa chọn, hãy đặt cùng Group Name cho tất cả radio của câu hỏi đó. "
        "Hai câu hỏi khác nhau phải dùng hai Group Name khác nhau.",
    )

    add_heading(document, "6. Liên kết", 1, page_break=True)
    add_heading(document, "6.1 Chèn liên kết vào chữ", 2)
    add_numbered_steps(
        document,
        [
            "Bôi đen chữ sẽ trở thành liên kết.",
            "Bấm nút Link trên toolbar/bubble menu hoặc chọn Insert → Link.",
            "Dán hoặc nhập địa chỉ, ví dụ https://example.com.",
            "Xác nhận để áp dụng. Sau khi chèn, việc gõ tiếp sẽ tiếp tục bên ngoài liên kết.",
        ],
    )
    add_heading(document, "6.2 Xem, sửa và gỡ liên kết", 2)
    add_bullets(
        document,
        [
            "Bấm vào phần chữ có liên kết để mở menu xem trước.",
            "Open Link mở liên kết trong tab mới.",
            "Edit Link cho phép sửa địa chỉ.",
            "Remove Link/Unlink gỡ liên kết nhưng giữ nguyên chữ.",
        ],
    )
    add_heading(document, "6.3 Địa chỉ hợp lệ", 2)
    add_paragraph(
        document,
        "UEditor kiểm tra an toàn trước khi chèn hoặc mở liên kết. Nên dùng HTTPS đầy đủ. "
        "Hệ thống có thể chấp nhận địa chỉ nội bộ, đường dẫn tương đối, email (mailto:) hoặc số điện thoại (tel:) "
        "nếu đúng định dạng; các địa chỉ nguy hiểm như javascript: bị từ chối."
    )
    add_callout(
        document,
        "warning",
        "Ctrl/Cmd+K",
        "Menu có thể hiển thị Ctrl+K bên cạnh Link, nhưng mã UEditor hiện tại không đăng ký phím này như một shortcut bắt buộc. "
        "Để thao tác ổn định, hãy dùng nút Link hoặc menu Insert → Link.",
    )

    add_heading(document, "7. Hình ảnh", 1, page_break=True)
    add_heading(document, "7.1 Các cách chèn ảnh", 2)
    add_table(
        document,
        ["Cách", "Thao tác", "Lưu ý"],
        [
            ["Từ URL", "Image → Add from URL → nhập URL ảnh và mô tả.", "URL phải hợp lệ và truy cập được."],
            ["Từ thiết bị", "Image → Upload → chọn một hoặc nhiều ảnh.", "Chỉ ảnh đúng định dạng/dung lượng cấu hình mới được chèn."],
            ["Dán", "Sao chép ảnh rồi Ctrl/Cmd+V trong editor.", "Ảnh được đặt tại vị trí con trỏ."],
            ["Kéo thả", "Kéo tệp ảnh vào vùng soạn thảo.", "Ảnh được đặt gần vị trí thả."],
        ],
        [3.2, 7.7, 5.8],
    )
    add_paragraph(
        document,
        "Các định dạng mặc định được hỗ trợ gồm PNG, JPEG, WebP, GIF và SVG; giới hạn mặc định là 10 MB mỗi ảnh. "
        "Màn hình nghiệp vụ có thể đặt giới hạn nhỏ hơn hoặc bỏ bớt định dạng."
    )
    add_heading(document, "7.2 Mô tả ảnh (Alt text)", 2)
    add_paragraph(
        document,
        "Khi chèn ảnh từ URL, nên điền Alt text ngắn gọn mô tả nội dung hoặc mục đích của ảnh. "
        "Với ảnh trang trí không mang thông tin, có thể để trống. Alt text giúp người dùng trình đọc màn hình và hiển thị khi ảnh lỗi."
    )
    add_heading(document, "7.3 Chọn bố cục và kích thước", 2)
    add_numbered_steps(
        document,
        [
            "Bấm vào ảnh để chọn. Ảnh được chọn có viền màu và xuất hiện menu nổi.",
            "Chọn Block để ảnh nằm trên một dòng riêng.",
            "Chọn Left hoặc Right để ảnh nổi bên trái/phải và chữ chạy bao quanh.",
            "Chọn S, M hoặc L để áp dụng kích thước định sẵn.",
            "Dùng Reset để trả về kích thước tự nhiên hoặc Delete để xóa ảnh.",
        ],
    )
    add_image(
        document,
        "06-anh-va-bubble-menu.png",
        "Hình 3. Ảnh được chọn, menu bố cục Block/Left/Right, cỡ S/M/L, Reset và Delete (giao diện tiếng Anh).",
        16.5,
    )
    add_heading(document, "7.4 Kéo đổi kích thước ảnh", 2)
    add_numbered_steps(
        document,
        [
            "Bấm hoặc rê chuột lên ảnh để hiện tay nắm nhỏ ở góc dưới bên phải.",
            "Giữ chuột trái tại tay nắm và kéo.",
            "Thả chuột khi đạt kích thước mong muốn. UEditor giữ tỷ lệ ảnh và không cho ảnh vượt quá chiều rộng editor.",
        ],
    )
    add_heading(document, "7.5 Ảnh khi lưu", 2)
    add_paragraph(
        document,
        "Tùy cấu hình, ảnh có thể được tải lên ngay khi chèn hoặc được giữ tạm trong nội dung rồi tải lên khi Lưu. "
        "Nếu màn hình dùng cơ chế tải lên khi lưu, không đóng trang cho đến khi có thông báo thành công."
    )
    add_callout(
        document,
        "danger",
        "Không đóng trang khi đang tải ảnh",
        "Nếu mất mạng hoặc đóng trang trong lúc ảnh/tệp đang tải, nội dung đã lưu có thể thiếu media. "
        "Hãy thử lại, kiểm tra ảnh hiển thị và chỉ rời trang sau khi hệ thống xác nhận lưu thành công.",
    )

    add_heading(document, "8. Lệnh nhanh bằng dấu /", 1, page_break=True)
    add_heading(document, "8.1 Cách mở và điều khiển", 2)
    add_numbered_steps(
        document,
        [
            "Đặt con trỏ tại một đoạn và gõ dấu /.",
            "Nhập tiếp từ khóa để lọc, ví dụ /table, /quote, /file.",
            "Dùng Arrow Up/Arrow Down để di chuyển giữa các kết quả.",
            "Nhấn Enter để chèn mục đang chọn; nhấn Esc để đóng danh sách.",
        ],
    )
    add_image(
        document,
        "03-lenh-gach-cheo.png",
        "Hình 4. Danh sách Basic Blocks sau khi gõ dấu / (giao diện tiếng Anh).",
        16.5,
    )
    add_heading(document, "8.2 Danh sách lệnh", 2)
    add_table(
        document,
        ["Lệnh/khối", "Kết quả", "Cách dùng chính"],
        [
            ["Text", "Đoạn văn thường.", "Chuyển đoạn hiện tại về văn bản."],
            ["Heading 1/2/3", "Tiêu đề cấp 1–3.", "Tổ chức cấu trúc nội dung."],
            ["Bullet List", "Danh sách dấu chấm.", "Liệt kê không thứ tự."],
            ["Numbered List", "Danh sách số.", "Quy trình hoặc thứ tự."],
            ["To-do List", "Danh sách công việc.", "Theo dõi trạng thái hoàn thành."],
            ["Quote", "Khối trích dẫn.", "Làm nổi bật lời dẫn."],
            ["Code Block", "Khối mã có tô màu.", "Chọn ngôn ngữ và sao chép mã."],
            ["Divider", "Đường kẻ ngang.", "Phân tách các phần."],
            ["Table", "Bảng 3×3 có hàng tiêu đề.", "Chèn nhanh bảng mặc định."],
            ["Callout", "Hộp thông tin có emoji.", "Làm nổi bật ghi chú quan trọng."],
            ["Bookmark Card", "Thẻ xem trước liên kết.", "Nhập URL trong hộp hỏi."],
            ["File Attachment", "Thẻ tệp đính kèm.", "Chọn tệp trên thiết bị."],
            ["Form Checkbox", "Checkbox hình vuông.", "Trường chọn độc lập."],
            ["Round Checkbox", "Checkbox hình tròn.", "Trường chọn độc lập kiểu tròn."],
            ["Form Radio Button", "Nút chọn một.", "Dùng Group Name để gom nhóm."],
        ],
        [4.2, 5.2, 7.3],
        compact=True,
    )

    add_heading(document, "9. Emoji và các khối nội dung nâng cao", 1, page_break=True)
    add_heading(document, "9.1 Chèn emoji bằng dấu :", 2)
    add_numbered_steps(
        document,
        [
            "Gõ dấu : và nhập tên emoji bằng tiếng Anh, ví dụ :heart.",
            "Dùng các phím mũi tên để di chuyển trong danh sách.",
            "Nhấn Enter để chèn emoji; nhấn Esc để đóng.",
            "Nếu không có kết quả, thử từ khóa tiếng Anh khác hoặc chỉ gõ : để xem emoji phổ biến.",
        ],
    )
    add_heading(document, "9.2 Hộp thông tin (Callout)", 2)
    add_numbered_steps(
        document,
        [
            "Gõ /callout và nhấn Enter.",
            "Bấm biểu tượng emoji ở đầu hộp để mở bộ chọn.",
            "Chọn emoji phù hợp.",
            "Chọn màu nền Gray, Blue, Green, Yellow hoặc Red.",
            "Bấm vào phần chữ của hộp và nhập nội dung như đoạn bình thường.",
        ],
    )
    add_heading(document, "9.3 Thẻ liên kết (Bookmark Card)", 2)
    add_paragraph(
        document,
        "Gõ /bookmark, nhập URL trong hộp hỏi và xác nhận. UEditor hiển thị tiêu đề, mô tả, tên miền và ảnh xem trước nếu hệ thống có dịch vụ lấy metadata. "
        "Bấm thẻ hoặc nhấn Enter khi thẻ đang được focus để mở liên kết. Nếu tải xem trước thất bại, dùng Retry Preview."
    )
    add_heading(document, "9.4 Tệp đính kèm (File Attachment)", 2)
    add_numbered_steps(
        document,
        [
            "Gõ /file và nhấn Enter.",
            "Chọn tệp trên thiết bị.",
            "Theo dõi trạng thái Uploading nếu hệ thống tải tệp lên ngay.",
            "Nếu có Upload failed, bấm Retry Upload.",
            "Bấm vào thẻ tệp hoặc nút Download để tải xuống.",
        ],
    )
    add_heading(document, "9.5 Khối mã (Code Block)", 2)
    add_bullets(
        document,
        [
            "Chọn ngôn ngữ: Plain Text, JavaScript, TypeScript, HTML, CSS, Python, Go, Rust, C++, Java, SQL, JSON, YAML hoặc Shell.",
            "Màu cú pháp được cập nhật theo ngôn ngữ.",
            "Rê chuột lên khối mã để hiện nút chọn ngôn ngữ và Copy.",
            "Nhấn Copy để sao chép toàn bộ mã; biểu tượng dấu kiểm báo thao tác thành công.",
            "Nhấn Enter ba lần liên tiếp ở cuối khối để thoát khối mã; hoặc dùng phím mũi tên xuống khi ở cuối khối.",
        ],
    )
    add_heading(document, "9.6 Đường phân cách", 2)
    add_paragraph(
        document,
        "Dùng /divider hoặc Insert → Horizontal Rule để chèn đường ngang. "
        "Đường phân cách phù hợp giữa các phần lớn, nhưng không nên lạm dụng thay cho tiêu đề."
    )

    add_heading(document, "10. Menu cổ điển", 1, page_break=True)
    add_image(
        document,
        "05-menu-tap-tin.png",
        "Hình 5. Menu File với Save và Export HTML (giao diện tiếng Anh).",
        16.5,
    )
    add_heading(document, "10.1 File", 2)
    add_table(
        document,
        ["Mục", "Tác dụng", "Điều kiện/lưu ý"],
        [
            ["Save", "Gọi hành động lưu của màn hình tích hợp.", "Chỉ dùng được khi hệ thống đã kết nối hàm lưu; menu có thể bị khóa."],
            ["Export HTML", "Xuất nội dung thành tệp document.html.", "Nếu hệ thống không cung cấp xử lý riêng, trình duyệt tự tải tệp HTML."],
        ],
        [3.6, 6.4, 6.7],
    )
    add_heading(document, "10.2 Edit", 2)
    add_bullets(
        document,
        [
            "Undo / Redo: hoàn tác và làm lại.",
            "Cut / Copy / Paste: thao tác clipboard.",
            "Paste as Text: dán nội dung không giữ định dạng.",
            "Select All: chọn toàn bộ tài liệu.",
        ],
    )
    add_heading(document, "10.3 View", 2)
    add_bullets(
        document,
        [
            "Source Code: mở hộp sửa HTML trực tiếp; bấm Apply để thay nội dung editor.",
            "Preview: xem nội dung ở chế độ đọc trước khi lưu.",
            "Fullscreen: đưa khung editor vào toàn màn hình nếu trình duyệt cho phép; thoát bằng Esc.",
            "Nút hình con mắt bên phải menu bar cũng mở Preview.",
        ],
    )
    add_callout(
        document,
        "danger",
        "Chỉnh sửa mã nguồn",
        "Source Code dành cho người hiểu HTML. HTML sai có thể làm hỏng bố cục hoặc mất thuộc tính của bảng, ảnh, công thức và trường tương tác. "
        "Nên sao chép nội dung dự phòng trước khi sửa.",
    )
    add_heading(document, "10.4 Insert", 2)
    add_bullets(
        document,
        [
            "Image / Image Upload: chèn ảnh từ URL hoặc thiết bị.",
            "Link: gắn liên kết vào chữ đã chọn.",
            "Table: chọn số hàng/cột bằng lưới tối đa 8×8.",
            "Form Fields: Form Checkbox, Round Checkbox và Form Radio Button.",
            "Horizontal Rule: chèn đường phân cách.",
        ],
    )
    add_heading(document, "10.5 Format", 2)
    add_bullets(
        document,
        [
            "Bold, Italic, Underline, Strikethrough, Superscript, Subscript, Inline Code.",
            "Wrap: Paragraph, Heading 1–3, Blockquote, Code Block, Task List.",
            "Align: Left, Center, Right, Justify.",
            "Clear Formatting: xóa định dạng.",
        ],
    )
    add_heading(document, "10.6 Tools và Table", 2)
    add_bullets(
        document,
        [
            "Tools → Source Code: mở trình chỉnh sửa HTML.",
            "Table → Insert Table: chèn bảng bằng lưới.",
            "Table → Delete Table: xóa toàn bộ bảng đang chọn.",
            "Table → Row/Column: xóa hàng hoặc cột.",
            "Table → Cell: Merge Cells hoặc Split Cell.",
            "Table Properties hiện là mục không thao tác trực tiếp; dùng menu ba chấm và menu ô để chỉnh thuộc tính thực tế.",
        ],
    )

    add_heading(document, "11. Bảng: hướng dẫn đầy đủ", 1, page_break=True)
    add_paragraph(
        document,
        "Chương này mô tả toàn bộ thao tác Table đang có trong UEditor: tạo và dán bảng, chọn ô, "
        "quản lý hàng/cột, kéo đổi kích thước, định dạng ô, viền và công thức. "
        "Để kiểm soát độ phủ, cuối chương có ma trận 131 thao tác và tùy chọn, đánh mã từ T001 đến T131."
    )
    add_callout(
        document,
        "info",
        "Cách hiểu con số 131",
        "Đây là số mục thao tác và tùy chọn người dùng có thể quan sát, bao gồm cả lựa chọn con như 9 vị trí viền, "
        "4 kiểu viền, 5 định dạng kết quả và 5 loại lỗi công thức. Đây không phải 131 mô-đun độc lập.",
    )

    add_heading(document, "11.1 Nhận biết các vùng điều khiển bảng", 2)
    add_image(
        document,
        "04-dieu-khien-bang.png",
        "Hình 6. Ô đang chọn, tay nắm hàng/cột, menu ba chấm và nút đổi kích thước bảng (giao diện tiếng Anh).",
        13.5,
    )
    add_table(
        document,
        ["Vùng điều khiển", "Vị trí", "Dùng để làm gì"],
        [
            ["Toolbar", "Phía trên vùng soạn thảo", "Chèn bảng; khi con trỏ ở trong bảng sẽ có thêm cột/hàng và gộp/tách ô."],
            ["Menu ba chấm", "Góc trên bên trái của bảng", "Căn cả bảng, thêm/xóa cấu trúc, bật hàng/cột tiêu đề và mở định dạng ô."],
            ["Tay nắm hàng", "Ngoài mép trái từng hàng", "Kéo đổi thứ tự hoặc mở menu riêng của hàng."],
            ["Tay nắm cột", "Phía trên từng cột", "Kéo đổi thứ tự hoặc mở menu riêng của cột."],
            ["Thanh dấu +", "Mép phải và mép dưới bảng", "Thêm nhanh một hoặc nhiều cột/hàng."],
            ["Nút kéo chéo", "Góc dưới bên phải", "Đổi kích thước toàn bộ bảng."],
            ["Thanh định dạng ô", "Nổi phía trên ô đang chọn", "Màu nền, viền, công thức, gộp/tách, wrap, căn dọc và hướng chữ."],
            ["Thanh công thức", "Phía trên nội dung editor", "Sửa công thức, chuyển thành giá trị hoặc xóa ô công thức."],
        ],
        [3.7, 4.3, 8.7],
        compact=True,
    )

    add_heading(document, "11.2 Chèn bảng mới", 2)
    add_image(
        document,
        "08-chen-bang-bang-luoi.png",
        "Hình 7. Chọn kích thước bảng trên lưới Insert Table (giao diện tiếng Anh).",
        16.0,
    )
    add_table(
        document,
        ["Cách chèn", "Các bước", "Kết quả"],
        [
            ["Toolbar", "Bấm Insert Table → rê đến số hàng/cột mong muốn → bấm.", "Tạo bảng từ 1×1 đến 8×8, hàng đầu là header."],
            ["Menu Insert", "Chọn Insert → Table → chọn kích thước trên lưới.", "Giống thao tác trên toolbar."],
            ["Menu Table", "Chọn Table → Insert Table → chọn kích thước.", "Dùng khi giao diện đang bật menu cổ điển."],
            ["Lệnh nhanh", "Gõ /table → nhấn Enter.", "Tạo nhanh bảng 3×3 có hàng tiêu đề."],
        ],
        [3.2, 8.1, 5.4],
    )
    add_callout(
        document,
        "tip",
        "Bảng lớn hơn 8×8",
        "Tạo kích thước gần nhất trên lưới, sau đó dùng Add Row, Add Column hoặc kéo thanh dấu + để mở rộng. "
        "Lưới chèn chỉ giới hạn lựa chọn ban đầu, không giới hạn kích thước cuối của bảng.",
    )

    add_heading(document, "11.3 Dán bảng từ Excel, Google Sheets hoặc trang web", 2)
    add_numbered_steps(
        document,
        [
            "Tại nguồn, chọn vùng dữ liệu cần sao chép và nhấn Ctrl/Cmd+C.",
            "Đặt con trỏ tại một dòng trống trong UEditor.",
            "Nhấn Ctrl/Cmd+V. Nếu clipboard có bảng HTML, UEditor tạo bảng theo cấu trúc đó. Nếu dữ liệu là văn bản có cột phân cách bằng Tab, UEditor tạo bảng từ dữ liệu TSV.",
            "Kiểm tra lại số hàng/cột, ô gộp, màu nền, viền, độ rộng, chiều cao và định dạng chữ.",
            "Kiểm tra lại công thức. Công thức của Excel/Google Sheets không được hiểu như công thức UEditor trừ khi dữ liệu đã có thuộc tính công thức tương thích.",
        ],
    )
    add_table(
        document,
        ["Dữ liệu có thể được giữ", "Điều kiện"],
        [
            ["Hàng, cột, ô tiêu đề, colspan và rowspan", "Nguồn sao chép phải cung cấp cấu trúc bảng HTML hợp lệ."],
            ["Độ rộng cột và chiều cao hàng", "Nguồn có thuộc tính hoặc CSS kích thước đọc được."],
            ["Màu nền, màu chữ, viền và căn nội dung", "Nguồn đưa định dạng vào HTML/CSS của clipboard."],
            ["Đậm, nghiêng, gạch chân, gạch ngang và nội dung nhiều dòng", "Nguồn cung cấp các thẻ/kiểu chữ tương ứng."],
            ["Công thức và định dạng số của UEditor", "Chỉ giữ khi HTML có các thuộc tính data-formula và data-number-format tương thích."],
        ],
        [7.0, 9.7],
        compact=True,
    )

    add_heading(document, "11.4 Chọn ô và di chuyển bằng bàn phím", 2)
    add_table(
        document,
        ["Thao tác", "Kết quả", "Khi nào nên dùng"],
        [
            ["Bấm một ô", "Đặt con trỏ để nhập hoặc sửa chữ.", "Soạn nội dung thông thường."],
            ["Nhấp đúp chuột trái lên phần chữ", "Bôi đen nội dung chữ trong ô.", "Định dạng chữ mà không định dạng cả ô."],
            ["Nhấp chuột trái ba lần lên phần chữ", "Chọn nguyên ô.", "Định dạng ô, mở công thức hoặc thực hiện thao tác theo ô."],
            ["Nhấp đúp chuột trái vào vùng trống trong ô", "Chọn nguyên ô.", "Dùng khi ô còn khoảng trống và không muốn bôi đen chữ."],
            ["Kéo qua nhiều ô", "Chọn vùng hình chữ nhật.", "Gộp ô hoặc áp dụng định dạng đồng thời."],
            ["Tab", "Sang ô kế tiếp.", "Nhập dữ liệu liên tục từ trái sang phải."],
            ["Shift+Tab", "Về ô trước.", "Sửa nhanh dữ liệu vừa nhập."],
            ["Tab tại ô cuối", "Tạo thêm một hàng và chuyển con trỏ vào hàng mới.", "Mở rộng bảng trong lúc nhập."],
            ["Backspace/Delete tại ô công thức đang chọn", "Xóa công thức và nội dung ô.", "Xóa nhanh ô tính toán."],
        ],
        [4.1, 7.0, 5.6],
        compact=True,
    )
    add_callout(
        document,
        "info",
        "Phân biệt bôi đen chữ và chọn ô",
        "Nhấp đúp chuột trái trực tiếp lên chữ để bôi đen phần chữ trong ô. "
        "Muốn chọn nguyên ô, nhấp chuột trái ba lần lên chữ hoặc nhấp đúp vào vùng trống trong ô.",
    )

    add_heading(document, "11.5 Menu ba chấm: lệnh áp dụng cho bảng", 2)
    add_image(
        document,
        "09-menu-dieu-khien-bang.png",
        "Hình 8. Menu điều khiển toàn bảng (giao diện tiếng Anh).",
        12.8,
    )
    add_table(
        document,
        ["Lệnh", "Tác dụng", "Điều kiện/lưu ý"],
        [
            ["Cell Formatting", "Chọn ô hiện tại để mở thanh định dạng ô.", "Cần bật menu nổi trên màn hình tích hợp."],
            ["Align Table Left", "Đặt toàn bộ bảng về bên trái.", "Không thay đổi căn chữ trong ô."],
            ["Align Table Center", "Đặt toàn bộ bảng ở giữa.", "Không thay đổi căn chữ trong ô."],
            ["Align Table Right", "Đặt toàn bộ bảng về bên phải.", "Không thay đổi căn chữ trong ô."],
            ["Add Column Before/After", "Thêm cột trước/sau cột hiện tại.", "Dữ liệu cũ được giữ."],
            ["Add Row Before/After", "Thêm hàng trên/dưới hàng hiện tại.", "Dữ liệu cũ được giữ."],
            ["Toggle Header Row", "Chuyển hàng đầu giữa ô tiêu đề và ô thường.", "Dùng cho tiêu đề theo hàng."],
            ["Toggle Header Column", "Chuyển cột đầu giữa ô tiêu đề và ô thường.", "Dùng cho tiêu đề theo cột."],
            ["Delete Column", "Xóa cột hiện tại.", "Không thể hoàn nguyên sau khi lưu; có thể dùng Undo trước khi lưu."],
            ["Delete Row", "Xóa hàng hiện tại.", "Công thức tham chiếu có thể đổi hoặc báo lỗi."],
            ["Delete Table", "Xóa toàn bộ bảng.", "Nên kiểm tra đúng bảng trước khi bấm."],
        ],
        [5.0, 6.4, 5.3],
        compact=True,
    )

    add_heading(document, "11.6 Menu hàng", 2)
    add_image(
        document,
        "10-menu-hang.png",
        "Hình 9. Menu của một hàng (giao diện tiếng Anh).",
        10.8,
    )
    add_table(
        document,
        ["Lệnh", "Cách thực hiện", "Kết quả"],
        [
            ["Add Row Before", "Rê mép trái hàng → mở tay nắm → chọn lệnh.", "Thêm hàng phía trên."],
            ["Add Row After", "Rê mép trái hàng → mở tay nắm → chọn lệnh.", "Thêm hàng phía dưới."],
            ["Duplicate Row", "Mở menu tay nắm hàng và chọn Duplicate Row.", "Sao chép nội dung và định dạng sang hàng mới phía dưới."],
            ["Clear Row Contents", "Mở menu tay nắm hàng và chọn Clear Row Contents.", "Làm trống nội dung nhưng giữ cấu trúc hàng."],
            ["Delete Row", "Mở menu tay nắm hàng và chọn Delete Row.", "Xóa cả hàng."],
            ["Move Row", "Giữ và kéo tay nắm hàng đến vị trí mới.", "Đổi thứ tự hàng."],
        ],
        [4.2, 7.4, 5.1],
        compact=True,
    )
    add_callout(
        document,
        "warning",
        "Hàng có công thức",
        "Duplicate Row sao chép cả thuộc tính của ô; Clear Row Contents giữ cấu trúc và định dạng ô. "
        "Sau hai thao tác này, hãy kiểm tra các ô công thức. Muốn xóa hẳn một công thức, dùng Clear Cell hoặc Backspace/Delete tại ô công thức.",
    )

    add_heading(document, "11.7 Menu cột", 2)
    add_image(
        document,
        "11-menu-cot.png",
        "Hình 10. Menu của một cột (giao diện tiếng Anh).",
        10.8,
    )
    add_table(
        document,
        ["Lệnh", "Cách thực hiện", "Kết quả"],
        [
            ["Add Column Before", "Rê phía trên cột → mở tay nắm → chọn lệnh.", "Thêm cột bên trái."],
            ["Add Column After", "Rê phía trên cột → mở tay nắm → chọn lệnh.", "Thêm cột bên phải."],
            ["Duplicate Column", "Mở menu tay nắm cột và chọn Duplicate Column.", "Sao chép nội dung, định dạng và độ rộng sang cột mới."],
            ["Clear Column Contents", "Mở menu tay nắm cột và chọn Clear Column Contents.", "Làm trống nội dung nhưng giữ cấu trúc cột."],
            ["Delete Column", "Mở menu tay nắm cột và chọn Delete Column.", "Xóa cả cột."],
            ["Move Column", "Giữ và kéo tay nắm cột đến vị trí mới.", "Đổi thứ tự cột."],
        ],
        [4.4, 7.2, 5.1],
        compact=True,
    )

    add_heading(document, "11.8 Thêm nhanh nhiều hàng hoặc cột", 2)
    add_numbered_steps(
        document,
        [
            "Bấm vào một ô để hiện các điều khiển bảng.",
            "Đưa chuột sát mép phải để hiện thanh dọc có dấu +. Bấm để thêm một cột; giữ và kéo sang phải để thêm nhiều cột.",
            "Đưa chuột sát mép dưới để hiện thanh ngang có dấu +. Bấm để thêm một hàng; giữ và kéo xuống để thêm nhiều hàng.",
            "Quan sát vùng xem trước rồi thả chuột để xác nhận số hàng/cột cần thêm.",
        ],
    )

    add_heading(document, "11.9 Đổi kích thước cột, hàng và toàn bảng", 2)
    add_table(
        document,
        ["Đối tượng", "Cách kéo", "Giới hạn và phím hỗ trợ"],
        [
            ["Cột", "Rê mép phải của ô đến khi hiện tay kéo màu xanh, giữ và kéo ngang.", "UEditor giữ độ rộng tối thiểu để cột còn sử dụng được."],
            ["Hàng", "Rê mép dưới của hàng đến khi hiện đường hướng dẫn, giữ và kéo dọc.", "Chiều cao tối thiểu được áp dụng tự động."],
            ["Toàn bộ bảng", "Kéo nút chéo ở góc dưới bên phải.", "Kéo tự do để đổi cả rộng và cao."],
            ["Khóa một trục", "Giữ Ctrl khi kéo toàn bảng.", "UEditor chọn trục thay đổi chính và giữ trục còn lại."],
            ["Giữ tỷ lệ", "Giữ Ctrl+Shift khi kéo toàn bảng.", "Tỷ lệ rộng/cao ban đầu được giữ."],
        ],
        [3.3, 7.6, 5.8],
    )
    add_callout(
        document,
        "tip",
        "Bảng rộng",
        "Nếu bảng rộng hơn vùng soạn thảo, dùng thanh cuộn ngang bên dưới bảng. "
        "Không nên giảm cột quá hẹp vì nội dung sẽ xuống dòng nhiều và làm bảng khó đọc.",
    )

    add_heading(document, "11.10 Hàng/cột tiêu đề và vị trí toàn bảng", 2)
    add_table(
        document,
        ["Mục tiêu", "Lệnh", "Gợi ý sử dụng"],
        [
            ["Hàng đầu là tiêu đề", "Toggle Header Row", "Dùng khi mỗi cột có tên dữ liệu."],
            ["Cột đầu là tiêu đề", "Toggle Header Column", "Dùng khi mỗi hàng có nhãn riêng."],
            ["Bảng nằm bên trái", "Align Table Left", "Phù hợp bảng hẹp trong nội dung dài."],
            ["Bảng nằm giữa", "Align Table Center", "Phù hợp bảng minh họa hoặc bảng tổng hợp."],
            ["Bảng nằm bên phải", "Align Table Right", "Chỉ dùng khi bố cục thực sự cần."],
        ],
        [4.2, 5.0, 7.5],
    )

    add_heading(document, "11.11 Gộp và tách ô", 2)
    add_numbered_steps(
        document,
        [
            "Kéo chọn ít nhất hai ô liền kề tạo thành một vùng hình chữ nhật.",
            "Bấm Merge Cells trên toolbar hoặc thanh định dạng ô.",
            "Để tách, chọn ô đã gộp rồi bấm Split Cell.",
            "Kiểm tra lại độ rộng cột, hàng/cột tiêu đề và công thức sau khi hoàn tất.",
        ],
    )
    add_callout(
        document,
        "warning",
        "Tọa độ ô có thể thay đổi",
        "Gộp, tách, thêm, xóa hoặc di chuyển hàng/cột có thể làm địa chỉ A1, B2… thay đổi. "
        "Luôn kiểm tra lại ô công thức và vùng tham chiếu trước khi lưu.",
    )

    add_heading(document, "11.12 Mở thanh định dạng ô", 2)
    add_image(
        document,
        "12-thanh-dinh-dang-o.png",
        "Hình 11. Thanh định dạng xuất hiện khi chọn ô (giao diện tiếng Anh).",
        12.8,
    )
    add_numbered_steps(
        document,
        [
            "Bấm một ô trong bảng.",
            "Mở menu ba chấm và chọn Cell Formatting; hoặc nhấp đúp vào vùng trống của ô.",
            "Thanh định dạng nổi xuất hiện phía trên ô. Từ trái sang phải gồm: Cell Background, Cell Border, Formula, Merge/Split, Wrap Text, Vertical Alignment và Text Direction.",
            "Nếu thanh không xuất hiện, dùng toolbar phía trên editor. Một số màn hình có thể tắt menu nổi.",
        ],
    )

    add_heading(document, "11.13 Màu nền và căn chữ theo chiều ngang", 2)
    add_table(
        document,
        ["Tính năng", "Cách dùng", "Lưu ý"],
        [
            ["Cell Background", "Chọn ô/vùng ô → bấm biểu tượng màu nền → chọn Automatic, màu có sẵn hoặc More Colors.", "Chọn Automatic để bỏ màu tùy chỉnh."],
            ["Align Left", "Chọn chữ hoặc ô → mở Alignment → Align Left.", "Chỉ căn nội dung, không di chuyển toàn bảng."],
            ["Align Center", "Chọn chữ hoặc ô → mở Alignment → Align Center.", "Phù hợp số liệu ngắn hoặc trạng thái."],
            ["Align Right", "Chọn chữ hoặc ô → mở Alignment → Align Right.", "Thường dùng cho số."],
            ["Justify", "Chọn nội dung → mở Alignment → Justify.", "Ít phù hợp với ô hẹp."],
        ],
        [3.4, 8.1, 5.2],
    )

    add_heading(document, "11.14 Căn dọc, hướng chữ và xuống dòng", 2)
    add_table(
        document,
        ["Nhóm", "Tùy chọn", "Hiệu quả"],
        [
            ["Vertical Alignment", "Top Align", "Đặt nội dung sát phía trên ô."],
            ["Vertical Alignment", "Middle Align", "Đặt nội dung ở giữa theo chiều cao."],
            ["Vertical Alignment", "Bottom Align", "Đặt nội dung sát phía dưới ô."],
            ["Text Direction", "Horizontal Text", "Hiển thị chữ theo chiều ngang thông thường."],
            ["Text Direction", "Vertical Text", "Xoay hướng trình bày chữ theo chiều dọc."],
            ["Text Wrap", "Wrap Text", "Cho phép nội dung tự xuống dòng trong độ rộng cột."],
            ["Text Wrap", "Don't Wrap Text", "Giữ nội dung trên một dòng; bảng/cột có thể cần cuộn ngang."],
        ],
        [4.2, 4.4, 8.1],
        compact=True,
    )

    add_heading(document, "11.15 Định dạng viền ô", 2)
    add_image(
        document,
        "13-dinh-dang-vien-o.png",
        "Hình 12. Bảng thiết lập viền ô (giao diện tiếng Anh).",
        13.5,
    )
    add_numbered_steps(
        document,
        [
            "Chọn một ô hoặc vùng nhiều ô.",
            "Mở Cell Formatting và bấm Cell Border.",
            "Chọn chế độ Draw Borders để vẽ hoặc Hide Borders để ẩn.",
            "Chọn vị trí viền, kiểu nét, độ dày và màu.",
            "Đóng bảng thiết lập rồi kiểm tra các cạnh chung giữa những ô liền kề.",
        ],
    )
    add_table(
        document,
        ["Nhóm", "Các lựa chọn", "Ý nghĩa"],
        [
            ["Vị trí", "All, Outer, Inner", "Tất cả; chỉ bao ngoài; hoặc chỉ các đường bên trong vùng chọn."],
            ["Vị trí bên trong", "Horizontal, Vertical", "Chỉ đường ngang hoặc chỉ đường dọc bên trong."],
            ["Từng cạnh", "Top, Right, Bottom, Left", "Chỉ áp dụng cạnh tương ứng của vùng chọn."],
            ["Xóa viền", "Clear Border", "Bỏ toàn bộ viền ở vị trí được áp dụng."],
            ["Chế độ", "Draw Borders, Hide Borders", "Vẽ theo kiểu đã chọn hoặc đặt viền thành ẩn."],
            ["Kiểu nét", "Solid, Dashed, Dotted, Double", "Liền, đứt, chấm hoặc hai đường."],
            ["Độ dày", "1px, 2px, 3px, 4px", "Độ dày đường viền."],
            ["Màu", "Bảng màu, More Colors", "Màu có sẵn hoặc màu tùy chọn của hệ thống."],
        ],
        [3.5, 6.0, 7.2],
        compact=True,
    )

    add_heading(document, "11.16 Công thức và định dạng số tại ô", 2)
    add_image(
        document,
        "14-menu-cong-thuc-o.png",
        "Hình 13. Bảng nhập công thức và chọn định dạng kết quả (giao diện tiếng Anh).",
        12.3,
    )
    add_table(
        document,
        ["Thành phần", "Cách dùng"],
        [
            ["Formula", "Nhập biểu thức, ví dụ =SUM(A2:A5). Nếu bỏ dấu =, UEditor tự bổ sung khi áp dụng từ bảng này."],
            ["Text", "Hiển thị giá trị thô."],
            ["123", "Định dạng Number có phân cách hàng nghìn."],
            ["$", "Định dạng Currency theo USD."],
            ["%", "Định dạng Percent."],
            ["Date", "Định dạng ngày từ số serial."],
            ["Apply", "Lưu công thức và tính kết quả."],
            ["Clear", "Xóa công thức khỏi ô."],
            ["Recalculate", "Tính lại công thức trong bảng đang chọn."],
        ],
        [4.1, 12.6],
        compact=True,
    )
    add_paragraph(
        document,
        "Chương 12 giải thích chi tiết cú pháp, hàm, tham chiếu ô/vùng, thanh công thức, định dạng kết quả và cách xử lý lỗi."
    )

    add_heading(document, "11.17 Phân biệt Clear, Delete và Convert to Value", 2)
    add_table(
        document,
        ["Lệnh", "Cái bị thay đổi", "Cái được giữ", "Dùng khi"],
        [
            ["Clear Row/Column Contents", "Nội dung hiển thị trong hàng/cột.", "Cấu trúc và định dạng ô.", "Muốn làm trống dữ liệu nhưng giữ bố cục."],
            ["Clear Cell", "Công thức và nội dung ô.", "Ô và vị trí trong bảng.", "Muốn xóa hẳn một ô công thức."],
            ["Delete Row/Column", "Dữ liệu và cấu trúc hàng/cột.", "Các phần còn lại của bảng.", "Không còn cần hàng/cột đó."],
            ["Delete Table", "Toàn bộ bảng.", "Nội dung ngoài bảng.", "Không còn cần bảng."],
            ["Convert to Value", "Thuộc tính công thức.", "Kết quả hiện tại dưới dạng dữ liệu cố định.", "Muốn khóa kết quả, không tự tính lại."],
        ],
        [3.6, 4.6, 4.4, 4.1],
        compact=True,
    )

    add_heading(document, "11.18 Quy trình thao tác bảng an toàn", 2)
    add_numbered_steps(
        document,
        [
            "Hoàn thiện cấu trúc hàng/cột trước khi nhập nhiều công thức.",
            "Bật hàng/cột tiêu đề và đặt tên rõ ràng.",
            "Nhập hoặc dán dữ liệu; kiểm tra ô gộp và định dạng.",
            "Điều chỉnh độ rộng cột, chiều cao hàng và vị trí toàn bảng.",
            "Áp dụng màu nền, viền, căn chữ và wrap.",
            "Tạo công thức sau cùng; kiểm tra kết quả và lỗi.",
            "Dùng Preview để kiểm tra bảng ở kích thước màn hình khách hàng sử dụng.",
            "Lưu và mở lại nội dung để xác nhận công thức, kích thước và định dạng còn nguyên.",
        ],
    )

    add_heading(document, "11.19 Những chức năng chưa có trong phiên bản hiện tại", 2)
    add_table(
        document,
        ["Chức năng", "Trạng thái", "Cách xử lý thay thế"],
        [
            ["Table Properties trong menu cổ điển", "Có nhãn nhưng đang bị khóa.", "Dùng menu ba chấm, thanh định dạng ô và tay nắm resize."],
            ["Sắp xếp và lọc dữ liệu", "Chưa có lệnh Sort/Filter.", "Sắp xếp dữ liệu ở nguồn trước khi dán."],
            ["Cố định hàng/cột", "Chưa có Freeze Row/Column.", "Giữ bảng ngắn hoặc chia thành nhiều bảng."],
            ["Định dạng có điều kiện", "Chưa có.", "Tô màu ô thủ công."],
            ["Biểu đồ từ dữ liệu bảng", "Chưa có.", "Tạo biểu đồ ngoài UEditor rồi chèn dưới dạng ảnh."],
            ["Chú thích ô và xác thực dữ liệu", "Chưa có.", "Dùng văn bản hướng dẫn hoặc Callout bên cạnh bảng."],
            ["Công thức giữa nhiều bảng/sheet", "Chưa hỗ trợ.", "Mọi tham chiếu chỉ dùng trong cùng một bảng."],
            ["Toàn bộ hàm Excel", "Chỉ có SUM, AVG, MIN, MAX và COUNT.", "Tính bên ngoài rồi dán kết quả nếu cần hàm khác."],
        ],
        [5.1, 4.2, 7.4],
        compact=True,
    )

    add_heading(document, "11.20 Ma trận kiểm tra 131 thao tác và tùy chọn Table", 2, page_break=True)
    add_paragraph(
        document,
        "Ma trận dưới đây dùng để tra cứu và kiểm tra độ phủ. Mỗi mã tương ứng một thao tác hoặc một lựa chọn người dùng có thể thực hiện trong nhóm Table."
    )
    add_table(
        document,
        ["Mã", "Nhóm", "Thao tác/tùy chọn", "Cách dùng ngắn gọn"],
        get_table_feature_rows(),
        [1.35, 3.0, 4.65, 7.7],
        compact=True,
    )

    add_heading(document, "12. Công thức trong bảng", 1, page_break=True)
    add_heading(document, "12.1 Khái niệm địa chỉ ô", 2)
    add_paragraph(
        document,
        "UEditor đánh địa chỉ giống bảng tính: cột dùng chữ A, B, C… và hàng dùng số 1, 2, 3… "
        "Ví dụ B2 là ô ở cột B, hàng 2. Vùng A2:A5 gồm các ô từ A2 đến A5 trong cùng bảng."
    )
    add_callout(
        document,
        "info",
        "Phạm vi công thức",
        "Tham chiếu chỉ hoạt động trong cùng một bảng. UEditor chưa hỗ trợ tên sheet, tham chiếu sang bảng khác hoặc cú pháp kiểu Sheet1!A1.",
    )
    add_heading(document, "12.2 Nhập công thức trực tiếp", 2)
    add_numbered_steps(
        document,
        [
            "Chọn ô sẽ hiển thị kết quả.",
            "Gõ dấu =. Danh sách hàm sẽ xuất hiện.",
            "Nhập tên hàm; dùng Arrow Up/Arrow Down để chọn và Tab để chèn mẫu hàm.",
            "Nhập địa chỉ ô/vùng, ví dụ =SUM(A2:A5). Có thể bấm hoặc kéo qua ô trong cùng bảng để chèn tham chiếu.",
            "Nhấn Enter để áp dụng và tính kết quả.",
            "Nhấn Escape trong lúc đang nhập để hủy công thức đang soạn.",
        ],
    )
    add_code(document, "=SUM(A2:A5)")
    add_code(document, "=(B2+C2)*10")
    add_code(document, "=MAX(A2:A10, C2)")
    add_heading(document, "12.3 Các hàm hỗ trợ", 2)
    add_table(
        document,
        ["Hàm", "Cú pháp ví dụ", "Ý nghĩa"],
        [
            ["SUM", "=SUM(A1:A5)", "Cộng các giá trị số; bỏ qua ô trống hoặc chữ trong vùng."],
            ["AVG", "=AVG(A1:A5)", "Tính trung bình các giá trị số; báo lỗi nếu không có số."],
            ["MIN", "=MIN(A1:A5)", "Giá trị nhỏ nhất."],
            ["MAX", "=MAX(A1:A5)", "Giá trị lớn nhất."],
            ["COUNT", "=COUNT(A1:A5)", "Đếm số ô có thể đọc như số."],
        ],
        [2.6, 4.8, 9.3],
    )
    add_heading(document, "12.4 Toán tử và cú pháp", 2)
    add_bullets(
        document,
        [
            "Hỗ trợ cộng (+), trừ (-), nhân (*) và chia (/).",
            "Hỗ trợ dấu ngoặc tròn để xác định thứ tự tính.",
            "Hỗ trợ số âm, số dương và số thập phân, ví dụ -1.5 hoặc .25.",
            "Nhiều đối số trong hàm được phân tách bằng dấu phẩy, ví dụ =MAX(A1:A3,B2).",
            "Giá trị có dấu phẩy hàng nghìn, ký hiệu $ ở đầu hoặc dấu % ở cuối có thể được đọc như số trong phép tính.",
            "Một tham chiếu trực tiếp đến ô chứa chữ sẽ báo lỗi; trong hàm tổng hợp, ô chữ/trống được bỏ qua.",
        ],
    )
    add_heading(document, "12.5 Chọn vùng bằng chuột", 2)
    add_paragraph(
        document,
        "Khi đang nhập công thức trong một ô, bấm một ô khác để chèn địa chỉ của ô đó. "
        "Giữ chuột và kéo qua nhiều ô để tạo vùng, ví dụ A2:C5. UEditor đánh dấu vùng được chọn. "
        "Nếu tham chiếu tạo vòng lặp, vùng bị đánh dấu màu cảnh báo và không được chèn."
    )
    add_heading(document, "12.6 Thanh công thức", 2)
    add_image(
        document,
        "07-cong-thuc-bang.png",
        "Hình 14. Ô B2 có công thức =SUM(A2:A3), kết quả 30 và thanh công thức (giao diện tiếng Anh).",
        16.5,
    )
    add_table(
        document,
        ["Thành phần", "Ý nghĩa"],
        [
            ["B2", "Địa chỉ ô đang chọn."],
            ["Ô công thức", "Hiển thị/sửa công thức gốc."],
            ["Dấu kiểm", "Áp dụng nội dung đang sửa."],
            ["Nút #", "Chuyển công thức thành giá trị cố định; sau đó ô không tự tính lại."],
            ["Thùng rác", "Xóa công thức và nội dung ô."],
        ],
        [4.0, 12.7],
    )
    add_heading(document, "12.7 Định dạng kết quả", 2)
    add_table(
        document,
        ["Định dạng", "Hiển thị", "Lưu ý"],
        [
            ["Text", "Giá trị thô.", "Mặc định."],
            ["Number", "Số có phân cách hàng nghìn, tối đa 6 chữ số thập phân.", "Định dạng en-US."],
            ["Currency", "Tiền tệ USD, tối đa 2 chữ số thập phân.", "Hiện ký hiệu $."],
            ["Percent", "Nhân giá trị hiển thị theo dạng phần trăm.", "Ví dụ 0.25 → 25%."],
            ["Date", "Ngày theo số serial kiểu Excel.", "Mốc 30/12/1899; hiển thị theo en-US trong bản hiện tại."],
        ],
        [3.0, 7.4, 6.3],
    )
    add_heading(document, "12.8 Tính lại tự động và thủ công", 2)
    add_bullets(
        document,
        [
            "Khi giá trị nguồn thay đổi, UEditor tự tính lại các ô công thức bị ảnh hưởng.",
            "Khi mở nội dung đã lưu, công thức trong các bảng được tính lại.",
            "Trong menu công thức của ô, bấm Recalculate để yêu cầu tính lại toàn bảng đang chọn.",
            "Nếu đang gõ công thức dở, UEditor chờ đến khi rời ô hoặc hoàn tất trước khi tính.",
        ],
    )
    add_heading(document, "12.9 Các lỗi công thức", 2)
    add_table(
        document,
        ["Thông báo", "Nguyên nhân thường gặp", "Cách xử lý"],
        [
            ["#EMPTY", "Công thức trống.", "Nhập biểu thức hoặc xóa công thức."],
            ["#INVALID-REFERENCE", "Ô tham chiếu không tồn tại hoặc không thể đọc như số.", "Kiểm tra địa chỉ và dữ liệu nguồn."],
            ["#INVALID-FORMULA", "Sai cú pháp, hàm không hỗ trợ hoặc thiếu ngoặc.", "Kiểm tra dấu =, tên hàm, dấu phẩy và ngoặc."],
            ["#DIVISION-BY-ZERO", "Chia cho 0 hoặc AVG không có giá trị số.", "Kiểm tra mẫu số/vùng dữ liệu."],
            ["#CIRCULAR-REFERENCE", "Công thức tham chiếu trực tiếp/gián tiếp về chính nó.", "Loại bỏ vòng tham chiếu."],
        ],
        [4.2, 6.4, 6.1],
        compact=True,
    )

    add_heading(document, "13. Phím tắt và thao tác bàn phím", 1, page_break=True)
    add_heading(document, "13.1 Phím tắt định dạng được UEditor xử lý", 2)
    add_table(
        document,
        ["Phím", "Tác dụng", "Ngữ cảnh"],
        [
            ["Ctrl/Cmd+B", "Bật/tắt in đậm.", "Chữ đã chọn hoặc chữ nhập tiếp."],
            ["Ctrl/Cmd+I", "Bật/tắt in nghiêng.", "Chữ đã chọn hoặc chữ nhập tiếp."],
            ["Ctrl/Cmd+U", "Bật/tắt gạch chân.", "Chữ đã chọn hoặc chữ nhập tiếp."],
            ["Ctrl/Cmd+Shift+S", "Bật/tắt gạch ngang.", "Chữ đã chọn hoặc chữ nhập tiếp."],
            ["Ctrl/Cmd+E", "Bật/tắt code nội dòng.", "Chữ đã chọn hoặc chữ nhập tiếp."],
            ["Ctrl/Cmd+,", "Bật/tắt chỉ số dưới.", "Chữ đã chọn hoặc chữ nhập tiếp."],
            ["Ctrl/Cmd+.", "Bật/tắt chỉ số trên.", "Chữ đã chọn hoặc chữ nhập tiếp."],
            ["Ctrl/Cmd+Alt+1/2/3", "Chuyển Heading 1/2/3.", "Đoạn hiện tại."],
            ["Ctrl/Cmd+Shift+8", "Danh sách dấu chấm.", "Đoạn/danh sách hiện tại."],
            ["Ctrl/Cmd+Shift+7", "Danh sách số.", "Đoạn/danh sách hiện tại."],
            ["Ctrl/Cmd+Shift+9", "Danh sách công việc.", "Đoạn/danh sách hiện tại."],
            ["Ctrl/Cmd+Shift+B", "Bật/tắt trích dẫn.", "Khối hiện tại."],
            ["Ctrl/Cmd+Alt+C", "Bật/tắt khối mã.", "Khối hiện tại."],
        ],
        [4.9, 6.2, 5.6],
        compact=True,
    )
    add_heading(document, "13.2 Lịch sử, chọn và clipboard", 2)
    add_table(
        document,
        ["Phím", "Tác dụng", "Ghi chú"],
        [
            ["Ctrl/Cmd+Z", "Hoàn tác.", "UEditor lưu tối đa khoảng 100 mốc lịch sử theo cấu hình hiện tại."],
            ["Ctrl/Cmd+Y", "Làm lại.", "Phổ biến trên Windows/Linux."],
            ["Shift+Cmd+Z", "Làm lại.", "Phổ biến trên macOS."],
            ["Ctrl/Cmd+A", "Chọn toàn bộ nội dung editor.", "Con trỏ phải ở trong editor."],
            ["Ctrl/Cmd+X", "Cắt.", "Do trình duyệt/hệ điều hành xử lý."],
            ["Ctrl/Cmd+C", "Sao chép.", "Do trình duyệt/hệ điều hành xử lý."],
            ["Ctrl/Cmd+V", "Dán.", "Có thể phụ thuộc quyền clipboard."],
        ],
        [4.4, 5.3, 7.0],
    )
    add_heading(document, "13.3 Phím theo ngữ cảnh", 2)
    add_table(
        document,
        ["Ngữ cảnh", "Phím", "Tác dụng"],
        [
            ["Lệnh /", "Arrow Up/Down", "Di chuyển kết quả."],
            ["Lệnh /", "Enter / Esc", "Chọn lệnh / đóng."],
            ["Emoji :", "Phím mũi tên", "Di chuyển trong lưới emoji."],
            ["Emoji :", "Enter / Esc", "Chèn emoji / đóng."],
            ["Gợi ý hàm", "Arrow Up/Down", "Di chuyển giữa SUM, AVG, MIN, MAX, COUNT."],
            ["Gợi ý hàm", "Tab", "Chèn hàm đang chọn và đặt con trỏ trong ngoặc."],
            ["Đang nhập công thức", "Enter / Escape", "Áp dụng/tính lại / hủy nhập."],
            ["Thanh công thức", "Enter / Escape", "Áp dụng bản nháp / bỏ thay đổi bản nháp."],
            ["Bảng", "Tab / Shift+Tab", "Ô kế tiếp / ô trước; Tab ở ô cuối thêm hàng."],
            ["Danh sách", "Tab / Shift+Tab", "Tăng / giảm cấp danh sách."],
            ["Đoạn văn", "Tab / Shift+Tab", "Chèn tab hoặc tăng / giảm lề tùy vị trí con trỏ."],
            ["Checkbox/Radio", "Space hoặc Enter", "Bật/tắt checkbox hoặc chọn radio."],
            ["Bookmark", "Enter", "Mở liên kết."],
            ["Menu nổi", "Escape", "Đóng menu."],
            ["Fullscreen", "Escape", "Thoát toàn màn hình theo trình duyệt."],
        ],
        [4.4, 4.5, 7.8],
        compact=True,
    )
    add_heading(document, "13.4 Những nhãn phím tắt cần hiểu đúng", 2)
    add_callout(
        document,
        "warning",
        "Save, Link và Fullscreen",
        "Menu hiển thị Ctrl+S cho Save, Ctrl+K cho Link và Ctrl+Shift+F cho Fullscreen. "
        "Trong UEditor hiện tại, đây là nhãn hướng dẫn trên menu chứ không phải cả ba đều được editor đăng ký trực tiếp. "
        "Khả năng hoạt động tùy màn hình tích hợp; luôn có thể dùng menu/nút tương ứng.",
    )

    add_heading(document, "14. Gõ nhanh theo cú pháp", 1, page_break=True)
    add_paragraph(
        document,
        "Ngoài phím tắt, UEditor nhận một số mẫu gõ kiểu Markdown. Nhập mẫu hoàn chỉnh rồi gõ khoảng trắng hoặc ký tự kết thúc theo mô tả."
    )
    add_table(
        document,
        ["Mẫu gõ", "Kết quả", "Ví dụ"],
        [
            ["# + Space", "Heading 1", "# Tiêu đề"],
            ["## + Space", "Heading 2", "## Tiêu đề"],
            ["### + Space", "Heading 3", "### Tiêu đề"],
            ["- + Space / * + Space / + + Space", "Bullet List", "- Mục"],
            ["1. + Space", "Numbered List", "1. Bước"],
            ["[ ] + Space hoặc - [ ] + Space", "Task List", "- [ ] Công việc"],
            ["> + Space", "Blockquote", "> Trích dẫn"],
            ["``` + Space/Enter", "Code Block", "```javascript"],
            ["---", "Horizontal Rule", "Gõ ba dấu gạch ngang."],
            ["**text** hoặc __text__", "Bold", "**quan trọng**"],
            ["*text* hoặc _text_", "Italic", "*nhấn mạnh*"],
            ["~~text~~", "Strikethrough", "~~đã hủy~~"],
            ["`text`", "Inline Code", "`npm run build`"],
        ],
        [5.4, 5.0, 6.3],
        compact=True,
    )
    add_callout(
        document,
        "tip",
        "Hoàn tác chuyển đổi tự động",
        "Nếu UEditor tự chuyển mẫu gõ nhưng bạn muốn giữ nguyên ký tự, nhấn Ctrl/Cmd+Z ngay sau khi chuyển đổi.",
    )

    add_heading(document, "15. Lưu, xem trước và xuất nội dung", 1, page_break=True)
    add_heading(document, "15.1 Lưu nội dung", 2)
    add_paragraph(
        document,
        "UEditor cung cấp nội dung cho màn hình nghiệp vụ; nút Lưu thực tế có thể nằm ngoài editor. "
        "Trong chế độ menu cổ điển, File → Save chỉ hoạt động khi ứng dụng đã kết nối hành động lưu."
    )
    add_numbered_steps(
        document,
        [
            "Kiểm tra nội dung, liên kết, ảnh, tệp, bảng và kết quả công thức.",
            "Bấm nút Lưu của màn hình hoặc File → Save.",
            "Chờ hệ thống tải ảnh/tệp nhúng lên máy chủ nếu có.",
            "Xác nhận thông báo thành công.",
            "Mở lại nội dung hoặc dùng Preview để kiểm tra.",
        ],
    )
    add_heading(document, "15.2 Xem trước", 2)
    add_paragraph(
        document,
        "Chọn View → Preview hoặc bấm biểu tượng con mắt. Hộp xem trước dùng chính nội dung HTML của editor ở chế độ đọc. "
        "Nếu nội dung trống, hệ thống hiển thị thông báo không có nội dung."
    )
    add_heading(document, "15.3 Xuất HTML", 2)
    add_paragraph(
        document,
        "File → Export HTML tải tệp document.html khi màn hình không thay thế bằng xử lý riêng. "
        "Tệp chứa HTML nội dung; giao diện hiển thị cuối cùng vẫn có thể phụ thuộc CSS của nơi mở tệp."
    )
    add_heading(document, "15.4 Chỉnh HTML trực tiếp", 2)
    add_numbered_steps(
        document,
        [
            "Chọn View/Tools → Source Code.",
            "Sao chép HTML hiện tại để dự phòng.",
            "Sửa phần cần thiết.",
            "Bấm Apply để cập nhật lại editor; bấm Close để hủy.",
            "Kiểm tra Preview và lưu.",
        ],
    )

    add_heading(document, "16. Thực hành tốt khi soạn nội dung", 1, page_break=True)
    add_heading(document, "16.1 Cấu trúc và khả năng đọc", 2)
    add_bullets(
        document,
        [
            "Mỗi tài liệu nên có tiêu đề rõ ràng và hệ thống Heading 1–3 nhất quán.",
            "Mỗi đoạn nên tập trung vào một ý; ưu tiên danh sách cho nhiều mục song song.",
            "Không dùng toàn chữ in hoa cho đoạn dài.",
            "Dùng Callout cho thông tin quan trọng, không dùng cho mọi đoạn.",
            "Dùng bảng cho dữ liệu có quan hệ hàng/cột, không dùng bảng chỉ để căn bố cục.",
        ],
    )
    add_heading(document, "16.2 Khả năng tiếp cận", 2)
    add_bullets(
        document,
        [
            "Viết Alt text cho ảnh có ý nghĩa.",
            "Dùng chữ liên kết mô tả mục tiêu, tránh chỉ ghi “bấm vào đây”.",
            "Bật hàng/cột tiêu đề cho bảng dữ liệu.",
            "Đảm bảo màu chữ và màu nền đủ tương phản.",
            "Không chỉ dùng màu, emoji hoặc hình ảnh để truyền đạt thông tin bắt buộc.",
            "Kiểm tra nội dung bằng bàn phím khi biểu mẫu có checkbox/radio.",
        ],
    )
    add_heading(document, "16.3 Hiệu năng và độ ổn định", 2)
    add_bullets(
        document,
        [
            "Tối ưu kích thước ảnh trước khi tải lên.",
            "Tránh dán tài liệu quá lớn một lần; chia thành từng phần và kiểm tra.",
            "Không chèn ảnh base64 rất lớn nếu hệ thống mạng yếu.",
            "Với bảng lớn, hạn chế gộp ô phức tạp và công thức tham chiếu quá rộng.",
            "Lưu theo từng giai đoạn đối với nội dung dài.",
        ],
    )

    add_heading(document, "17. Xử lý sự cố thường gặp", 1, page_break=True)
    add_table(
        document,
        ["Hiện tượng", "Nguyên nhân có thể", "Cách xử lý"],
        [
            ["Không thấy toolbar/menu", "Màn hình dùng variant tối giản, chế độ chỉ đọc hoặc đã tắt chức năng.", "Kiểm tra quyền chỉnh sửa; cuộn ngang toolbar; liên hệ quản trị nếu cần."],
            ["Nút bị mờ", "Chưa chọn đúng đối tượng hoặc lệnh không hợp lệ.", "Đặt con trỏ/chọn chữ, ảnh hoặc ô phù hợp."],
            ["Không dán được", "Trình duyệt chặn clipboard.", "Dùng Ctrl/Cmd+V và cấp quyền clipboard nếu được hỏi."],
            ["Link bị từ chối", "URL thiếu/sai định dạng hoặc bị chặn vì an toàn.", "Dùng URL HTTPS đầy đủ và kiểm tra ký tự."],
            ["Ảnh không được chèn", "Sai định dạng, quá dung lượng hoặc upload lỗi.", "Giảm kích thước, đổi PNG/JPEG/WebP, kiểm tra mạng và thử lại."],
            ["Ảnh/tệp báo Upload failed", "Mạng hoặc dịch vụ upload lỗi.", "Bấm Retry; nếu vẫn lỗi, tải lại trang sau khi đã sao lưu nội dung."],
            ["Không thấy menu nổi", "Editor mất focus, chưa chọn đối tượng hoặc menu bị tắt.", "Bấm lại editor, bôi đen/chọn đối tượng; dùng toolbar thay thế."],
            ["Không thấy tay nắm bảng", "Chuột chưa ở vùng mép hoặc bảng ở chế độ chỉ đọc.", "Bấm trong bảng rồi rê phía trên, trái, phải hoặc góc dưới."],
            ["Không gộp được ô", "Vùng chọn không liền kề/hình chữ nhật.", "Chọn lại các ô liền nhau."],
            ["Công thức không tính", "Công thức đang nhập dở, sai cú pháp hoặc tham chiếu lỗi.", "Nhấn Enter, kiểm tra thông báo lỗi và dùng Recalculate."],
            ["Bảng thay đổi sau khi dán", "Nguồn có định dạng không tương thích hoàn toàn.", "Kiểm tra lại header, gộp ô, viền, công thức và độ rộng."],
            ["Save không bấm được", "Màn hình chưa tích hợp hành động Save trong menu.", "Dùng nút Lưu của màn hình nghiệp vụ."],
            ["Fullscreen/clipboard không hoạt động", "Trình duyệt từ chối quyền hoặc shortcut chưa được màn hình bắt.", "Dùng nút/menu và kiểm tra quyền trình duyệt."],
        ],
        [4.2, 6.1, 6.4],
        compact=True,
    )
    add_heading(document, "17.1 Quy trình phục hồi an toàn", 2)
    add_numbered_steps(
        document,
        [
            "Dừng thao tác gây lỗi; thử Ctrl/Cmd+Z.",
            "Sao chép phần nội dung quan trọng sang nơi tạm.",
            "Kiểm tra mạng nếu có ảnh/tệp.",
            "Lưu lại nếu hệ thống vẫn cho phép.",
            "Tải lại trang và kiểm tra nội dung đã lưu.",
            "Nếu lỗi lặp lại, ghi lại bước thao tác, trình duyệt, ảnh chụp và thời điểm để gửi hỗ trợ.",
        ],
    )

    add_heading(document, "18. Danh mục tính năng đầy đủ", 1, page_break=True)
    add_paragraph(
        document,
        "Tài liệu kiểm kê 231 mục thao tác và tùy chọn người dùng: 100 mục UEditor ngoài bảng (U001–U100) "
        "và 131 mục thuộc Table (T001–T131). Cách đếm này dùng để kiểm soát độ phủ; các lựa chọn con được tính riêng để dễ tra cứu."
    )
    add_heading(document, "18.1 Tổng quan theo nhóm", 2)
    add_table(
        document,
        ["Nhóm", "Tính năng UEditor"],
        [
            ["Soạn thảo", "Nhập HTML rich text, chọn, xóa, cắt/chép/dán, dán text, hoàn tác/làm lại, bộ đếm từ và ký tự, giới hạn ký tự."],
            ["Kiểu chữ", "Font family, cỡ 8–96 px, giãn dòng, giãn chữ, Normal, Heading 1–3."],
            ["Định dạng inline", "Bold, Italic, Underline, Strikethrough, Inline Code, Subscript, Superscript."],
            ["Màu và đoạn", "Màu chữ, highlight, căn trái/giữa/phải/đều, tăng/giảm lề."],
            ["Danh sách", "Bullet, numbered, task list, danh sách lồng nhau."],
            ["Liên kết", "Chèn, xem trước, mở, sửa, gỡ và kiểm tra URL an toàn."],
            ["Media", "Ảnh URL/upload/paste/drop, alt text, resize, block/left/right, S/M/L, reset, delete; tệp đính kèm và tải xuống."],
            ["Khối nội dung", "Quote, code block, divider, callout, bookmark, checkbox vuông/tròn, radio."],
            ["Emoji", "Gợi ý bằng dấu :, tìm kiếm, điều khiển bàn phím, emoji trong callout."],
            ["Lệnh nhanh", "15 khối qua dấu /, lọc và điều khiển bằng bàn phím."],
            ["Bảng cấu trúc", "Chèn qua lưới hoặc /table; dán HTML/TSV; thêm, xóa, nhân đôi, di chuyển và làm trống hàng/cột; bật hàng/cột tiêu đề."],
            ["Bảng kích thước", "Resize cột, hàng và toàn bảng; thêm nhanh nhiều hàng/cột; Ctrl khóa trục; Ctrl+Shift giữ tỷ lệ; cuộn ngang bảng rộng."],
            ["Bảng ô", "Chọn một/nhiều ô, merge/split, màu nền, 9 vị trí viền, 4 kiểu nét, 4 độ dày, màu tùy chọn, căn ngang/dọc, chữ ngang/dọc và wrap."],
            ["Công thức", "Tham chiếu ô/vùng, chọn vùng bằng chuột, + − × ÷, SUM/AVG/MIN/MAX/COUNT, tự tính lại, 5 lỗi và 5 định dạng kết quả."],
            ["Ma trận Table", "131 thao tác và tùy chọn có mã T001–T131 để kiểm tra độ phủ và tra cứu nhanh."],
            ["Menu cổ điển", "File, Edit, View, Insert, Format, Tools, Table; Preview, Source Code, Fullscreen, Export HTML."],
            ["Chế độ hiển thị", "Default, minimal, medium, medium-full, full, notion; chế độ chỉ đọc."],
        ],
        [4.0, 12.7],
        compact=True,
    )
    add_heading(document, "18.2 Ma trận 100 thao tác và tùy chọn ngoài Table", 2, page_break=True)
    add_paragraph(
        document,
        "Các mục Table không lặp lại trong bảng dưới đây; xem ma trận T001–T131 tại mục 11.20."
    )
    add_table(
        document,
        ["Mã", "Nhóm", "Thao tác/tùy chọn", "Cách dùng ngắn gọn"],
        get_core_feature_rows(),
        [1.35, 3.0, 4.65, 7.7],
        compact=True,
    )

    add_heading(document, "19. Thuật ngữ đối chiếu Anh – Việt", 1, page_break=True)
    add_table(
        document,
        ["English label", "Nghĩa dùng trong tài liệu", "English label", "Nghĩa dùng trong tài liệu"],
        [
            ["File", "Tập tin", "Edit", "Sửa"],
            ["View", "Xem", "Insert", "Chèn/Thêm"],
            ["Format", "Định dạng", "Tools", "Công cụ"],
            ["Table", "Bảng", "Save", "Lưu"],
            ["Export HTML", "Xuất HTML", "Source Code", "Mã nguồn"],
            ["Preview", "Xem trước", "Fullscreen", "Toàn màn hình"],
            ["Bold", "In đậm", "Italic", "In nghiêng"],
            ["Underline", "Gạch chân", "Strikethrough", "Gạch ngang"],
            ["Inline Code", "Mã nội dòng", "Code Block", "Khối mã"],
            ["Subscript", "Chỉ số dưới", "Superscript", "Chỉ số trên"],
            ["Text Color", "Màu chữ", "Highlight", "Màu đánh dấu"],
            ["Bullet List", "Danh sách dấu chấm", "Numbered List", "Danh sách số"],
            ["Task List", "Danh sách công việc", "Blockquote", "Trích dẫn"],
            ["Callout", "Hộp thông tin", "Bookmark Card", "Thẻ liên kết"],
            ["File Attachment", "Tệp đính kèm", "Form Fields", "Trường biểu mẫu"],
            ["Merge Cells", "Gộp ô", "Split Cell", "Tách ô"],
            ["Wrap Text", "Tự xuống dòng", "Text Direction", "Hướng chữ"],
            ["Recalculate", "Tính lại", "Convert to Value", "Chuyển thành giá trị"],
        ],
        [3.4, 5.0, 3.4, 5.0],
        compact=True,
    )

    return document


def iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested_table in cell.tables:
                yield from iter_table_paragraphs(nested_table)


def finalize_typefaces(document: Document) -> None:
    """Normalize every generated run so Word never falls back for Vietnamese."""
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        paragraphs.extend(iter_table_paragraphs(table))

    for section in document.sections:
        for part in (section.header, section.footer):
            paragraphs.extend(part.paragraphs)
            for table in part.tables:
                paragraphs.extend(iter_table_paragraphs(table))

    for paragraph in paragraphs:
        typeface = CODE_FONT if paragraph.style and paragraph.style.name == "Guide Code" else BODY_FONT
        for run in paragraph.runs:
            set_run_typeface(run, typeface)


def main() -> None:
    document = build_document()
    finalize_typefaces(document)
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
